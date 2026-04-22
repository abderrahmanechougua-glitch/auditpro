"""
Module : Circularisation des Tiers (NEP 505 / ISA 505)
Workflow en 4 étapes indépendantes — chaque étape a ses propres inputs.
"""
import os
import importlib.util
import io
from contextlib import redirect_stdout
from pathlib import Path
from datetime import datetime

import pandas as pd

from modules.base_module import BaseModule, ModuleInput, ModuleResult

# ── Définition des 4 étapes ────────────────────────────────────

ETAPES = [
    {
        "id":          "selection",
        "label":       "Étape 1 — Sélection des tiers",
        "description": "Charge une balance auxiliaire ou un GL et sélectionne les tiers à circulariser par cumul décroissant (85-95 %).\n\nSortie : fichier Excel des tiers sélectionnés (à compléter avec emails/adresses).",
        "icon":        "1",
    },
    {
        "id":          "lettres",
        "label":       "Étape 2 — Génération des lettres",
        "description": "Génère les lettres de circularisation Word individuelles à partir du fichier de sélection et d'un modèle Word (canvas).\n\nSortie : lettres .docx individuelles + document consolidé.",
        "icon":        "2",
    },
    {
        "id":          "split",
        "label":       "Étape 3 — Split PDF signé",
        "description": "Découpe le PDF signé (imprimé, signé, scanné puis remis en PDF) en autant de fichiers individuels que de tiers.\n\nSortie : PDF individuels nommés par tiers.",
        "icon":        "3",
    },
    {
        "id":          "emails",
        "label":       "Étape 4 — Génération des emails",
        "description": "Génère les fichiers .eml prêts à l'envoi, avec la lettre PDF en pièce jointe et le corps personnalisé par tiers.\n\nSortie : fichiers .eml + rapport des emails manquants.",
        "icon":        "4",
    },
]

# ── Inputs par étape ───────────────────────────────────────────

INPUTS_PAR_ETAPE = {

    "selection": [
        ModuleInput(
            key="balance_aux",
            label="Balance auxiliaire / Grand Livre",
            input_type="file",
            extensions=[".xlsx", ".xls", ".csv"],
            tooltip="Fichier Excel avec les soldes ou mouvements par tiers (clients ou fournisseurs)",
        ),
        ModuleInput(
            key="type_tiers",
            label="Type de tiers",
            input_type="combo",
            options=["Fournisseurs", "Clients"],
            default="Fournisseurs",
            required=True,
        ),
        ModuleInput(
            key="couverture_min",
            label="Couverture minimale (%)",
            input_type="number",
            default=85,
            required=False,
            tooltip="Ex : 85 → sélectionner au moins 85 % du montant total",
        ),
        ModuleInput(
            key="couverture_max",
            label="Couverture maximale (%)",
            input_type="number",
            default=95,
            required=False,
            tooltip="Ex : 95 → ne pas dépasser 95 % du montant total",
        ),
    ],

    "lettres": [
        ModuleInput(
            key="fichier_selection",
            label="Fichier de sélection (Excel)",
            input_type="file",
            extensions=[".xlsx", ".xls"],
            tooltip="Fichier produit par l'étape 1 (selection/output/)",
        ),
        ModuleInput(
            key="date_lettre",
            label="Date de la lettre",
            input_type="text",
            default=datetime.now().strftime("%d/%m/%Y"),
            tooltip="Date à afficher dans les lettres (JJ/MM/AAAA)",
        ),
    ],

    "split": [
        ModuleInput(
            key="pdf_signe",
            label="PDF signé à découper",
            input_type="file",
            extensions=[".pdf"],
            tooltip="PDF contenant toutes les lettres signées (une ou plusieurs pages par lettre)",
        ),
        ModuleInput(
            key="fichier_selection",
            label="Fichier de sélection (pour nommer les PDF)",
            input_type="file",
            extensions=[".xlsx", ".xls"],
            required=False,
            tooltip="Optionnel — utilisé pour nommer les PDF par tiers. Si absent, nommage automatique.",
        ),
        ModuleInput(
            key="pages_par_lettre",
            label="Pages par lettre",
            input_type="number",
            default=1,
            tooltip="Nombre de pages par lettre dans le PDF (habituellement 1 ou 2)",
        ),
    ],

    "emails": [
        ModuleInput(
            key="fichier_selection",
            label="Fichier de sélection avec emails",
            input_type="file",
            extensions=[".xlsx", ".xls"],
            tooltip="Fichier Excel avec les colonnes email et nom des tiers (produit étape 1, complété)",
        ),
        ModuleInput(
            key="dossier_pdf",
            label="Dossier des lettres PDF individuelles",
            input_type="folder",
            tooltip="Dossier contenant les PDF produits par l'étape 3 (lettres/output/)",
        ),
        ModuleInput(
            key="nom_societe",
            label="Nom de la société auditée",
            input_type="text",
        ),
        ModuleInput(
            key="date_arrete",
            label="Date d'arrêté",
            input_type="text",
            default="31/12/2024",
        ),
        ModuleInput(
            key="type_tiers",
            label="Type de tiers",
            input_type="combo",
            options=["Clients", "Fournisseurs"],
            default="Clients",
            required=False,
        ),
    ],
}


class CircularisationTiers(BaseModule):

    name = "Circularisation des Tiers"
    description = (
        "Automatisation NEP 505 / ISA 505 — 4 étapes indépendantes :\n"
        "Sélection → Lettres → Split PDF → Emails"
    )
    category = "Audit"
    version = "3.0"
    help_text = (
        "WORKFLOW EN 4 ÉTAPES INDÉPENDANTES\n\n"
        "1. SÉLECTION : charge la balance auxiliaire et sélectionne les tiers\n"
        "   à circulariser par cumul décroissant (85-95 %).\n\n"
        "2. GÉNÉRATION DES LETTRES : remplit un modèle Word avec les données\n"
        "   de chaque tiers (nom, montant, date de clôture).\n\n"
        "3. SPLIT PDF : découpe le PDF signé en fichiers PDF individuels\n"
        "   nommés par tiers, prêts à être envoyés.\n\n"
        "4. GÉNÉRATION EMAILS : crée des fichiers .eml personnalisés avec\n"
        "   le PDF en pièce jointe — à ouvrir dans Outlook pour envoi.\n\n"
        "Vous pouvez exécuter chaque étape séparément selon votre avancement."
    )
    detection_keywords = ["client", "fournisseur", "tiers", "circularisation",
                          "confirmation", "balance auxiliaire", "solde"]
    detection_threshold = 0.4

    # Module-level step definitions — used by the generic workspace step-selector
    ETAPES = ETAPES
    INPUTS_PAR_ETAPE = INPUTS_PAR_ETAPE

    # Étape active (modifiée par l'UI)
    etape_active: str = "selection"

    # ── API BaseModule ─────────────────────────────────────────

    def get_required_inputs(self):
        return INPUTS_PAR_ETAPE.get(self.etape_active, [])

    def get_param_schema(self):
        return []  # Pas de paramètres cachés — tout est dans les inputs

    def validate(self, inputs):
        errors = []
        etape = inputs.get("_etape", self.etape_active)
        required = INPUTS_PAR_ETAPE.get(etape, [])
        for inp in required:
            if not inp.required:
                continue
            val = inputs.get(inp.key, "")
            if not val:
                errors.append(f"'{inp.label}' est requis.")
            elif inp.input_type == "file" and not Path(str(val)).exists():
                errors.append(f"Fichier introuvable : {val}")
            elif inp.input_type == "folder" and not Path(str(val)).exists():
                errors.append(f"Dossier introuvable : {val}")
        return (not errors, errors)

    def preview(self, inputs):
        etape = inputs.get("_etape", self.etape_active)
        # Aperçu du fichier principal selon l'étape
        file_keys = {
            "selection": "balance_aux",
            "lettres":   "fichier_selection",
            "split":     "fichier_selection",
            "emails":    "fichier_selection",
        }
        key = file_keys.get(etape)
        if key:
            p = inputs.get(key, "")
            if p and Path(str(p)).exists() and str(p).endswith((".xlsx", ".xls")):
                try:
                    return pd.read_excel(p, nrows=10)
                except Exception:
                    pass
        return None

    def execute(self, inputs, output_dir, progress_callback=None):
        etape = inputs.get("_etape", self.etape_active)
        os.makedirs(output_dir, exist_ok=True)

        if progress_callback:
            progress_callback(5, f"Démarrage — {etape}...")

        # Charger le moteur script.py
        script_path = Path(__file__).parent / "script.py"
        if not script_path.exists():
            return ModuleResult(
                success=False,
                message="script.py introuvable dans modules/circularisation/\n"
                        "Ajoutez le fichier script.py du projet circularisation."
            )

        spec = importlib.util.spec_from_file_location("circ_engine", str(script_path))
        eng  = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(eng)

        # Surcharger SCRIPT_DIR pour que les sorties aillent dans output_dir
        eng.SCRIPT_DIR = Path(output_dir)
        eng.DIRS = {k: str(Path(output_dir) / v) for k, v in eng.DIRS.items()}

        # Charger la config
        config = eng.charger_config() if hasattr(eng, "charger_config") else eng.DEFAULT_CONFIG
        # Créer les dossiers nécessaires
        if hasattr(eng, "creer_dossiers"):
            # patch SCRIPT_DIR dans le module
            import types
            eng.SCRIPT_DIR = Path(output_dir)
            eng.creer_dossiers()

        captured = io.StringIO()

        try:
            with redirect_stdout(captured):
                if etape == "selection":
                    result_path = self._run_selection(eng, inputs, config, output_dir, progress_callback)
                elif etape == "lettres":
                    result_path = self._run_lettres(eng, inputs, config, output_dir, progress_callback)
                elif etape == "split":
                    result_path = self._run_split(eng, inputs, config, output_dir, progress_callback)
                elif etape == "emails":
                    result_path = self._run_emails(eng, inputs, config, output_dir, progress_callback)
                else:
                    return ModuleResult(success=False, message=f"Étape inconnue : {etape}")

            if progress_callback:
                progress_callback(100, "Terminé !")

            logs = [l.strip() for l in captured.getvalue().splitlines() if l.strip()]
            log_summary = "\n".join(f"• {l}" for l in logs[-6:]) if logs else ""

            rp = Path(str(result_path)) if result_path else None

            if rp and rp.exists():
                # ── Cas : résultat = DOSSIER (split PDF) ──────────
                if rp.is_dir():
                    pdfs = sorted(rp.glob("*.pdf"))
                    nb   = len(pdfs)
                    taille_totale = sum(f.stat().st_size for f in pdfs) // 1024
                    lignes = [f"   {p.name}  ({p.stat().st_size//1024} Ko)" for p in pdfs[:8]]
                    if nb > 8:
                        lignes.append(f"   … et {nb-8} autres")
                    msg = (
                        f"{nb} PDF créé(s) avec succès\n\n"
                        f"Dossier : {rp.name}\n"
                        f"Taille totale : {taille_totale} Ko\n\n"
                        + "\n".join(lignes)
                    )
                    return ModuleResult(
                        success=True,
                        output_path=str(rp),
                        message=msg,
                        stats={"PDFs créés": nb, "Taille (Ko)": taille_totale},
                    )

                # ── Cas : résultat = FICHIER (sélection, lettres, emails) ──
                ext  = rp.suffix.lower()
                size = rp.stat().st_size // 1024
                labels = {
                    ".xlsx": "Fichier Excel",
                    ".docx": "Document Word",
                    ".pdf":  "PDF",
                    ".eml":  "Email",
                }
                type_label = labels.get(ext, "Fichier")
                msg = (
                    f"{type_label} généré avec succès\n\n"
                    f"Fichier : {rp.name}\n"
                    f"Taille  : {size} Ko"
                    + (f"\n\n{log_summary}" if log_summary else "")
                )
                return ModuleResult(
                    success=True,
                    output_path=str(rp),
                    message=msg,
                )

            # ── Cas : aucun résultat explicite → chercher le plus récent ──
            all_items = [f for f in Path(output_dir).rglob("*") if f.is_file()]
            if all_items:
                newest = max(all_items, key=lambda f: f.stat().st_mtime)
                return ModuleResult(
                    success=True,
                    output_path=str(newest),
                    message=f"Traitement terminé.\n\nFichier : {newest.name}"
                            + (f"\n\n{log_summary}" if log_summary else ""),
                )
            return ModuleResult(
                success=False,
                message=f"Aucun fichier produit.\n\n{log_summary}"
            )

        except Exception as e:
            logs = [l.strip() for l in captured.getvalue().splitlines() if l.strip()]
            detail = "\n".join(f"• {l}" for l in logs[-6:])
            return ModuleResult(
                success=False,
                errors=[f"Erreur étape '{etape}' : {e}\n\n{detail}"]
            )

    # ── Runners par étape ──────────────────────────────────────

    def _run_selection(self, eng, inputs, config, output_dir, cb):
        if cb: cb(20, "Chargement du fichier…")
        chemin = inputs["balance_aux"]
        type_tiers = inputs.get("type_tiers", "Fournisseurs").lower()
        type_tiers = "fournisseurs" if "fourn" in type_tiers else "clients"

        seuil_min = float(inputs.get("couverture_min", 85)) / 100
        seuil_max = float(inputs.get("couverture_max", 95)) / 100

        if cb: cb(35, "Analyse du format…")

        if eng.est_format_dynamics(chemin):
            df, societe = eng.parser_dynamics_ax(chemin, type_tiers)
            col_montant = "total_credit" if type_tiers == "fournisseurs" else "solde_closing"
        else:
            df = eng.charger_fichier(chemin)
            if df is None or df.empty:
                raise ValueError("Fichier vide ou illisible.")
            mapping = _auto_mapper(df, type_tiers, eng, config)
            col_tiers   = mapping.get("tiers")
            col_nom     = mapping.get("nom_tiers", col_tiers)
            col_montant = mapping.get("montant_credit" if type_tiers == "fournisseurs" else "solde", None)
            if not col_tiers or not col_montant:
                raise ValueError(
                    f"Colonnes non identifiées automatiquement.\n"
                    f"Colonnes disponibles : {list(df.columns)}\n"
                    f"Renommez-les en : tiers, nom, solde (ou montant_credit)"
                )
            df[col_montant] = pd.to_numeric(df[col_montant], errors="coerce").fillna(0)
            df = df.groupby(col_tiers).agg(
                {col_montant: "sum", **({col_nom: "first"} if col_nom != col_tiers else {})}
            ).reset_index()
            df = df.rename(columns={col_tiers: "code_tiers", col_montant: "montant"})
            if col_nom and col_nom != col_tiers and col_nom in df.columns:
                df = df.rename(columns={col_nom: "nom"})
            else:
                df["nom"] = df["code_tiers"]
            col_montant = "montant"

        if cb: cb(55, "Sélection par cumul décroissant…")

        df_positifs = df[df[col_montant].abs() > 0]
        sel, stats  = eng.selectionner_par_cumul(df_positifs, col_montant, seuil_min, seuil_max)

        if "email" not in sel.columns:   sel["email"]   = ""
        if "adresse" not in sel.columns: sel["adresse"] = ""

        if cb: cb(75, "Export Excel…")

        date_str  = datetime.now().strftime("%Y%m%d_%H%M")
        out_dir   = Path(output_dir)
        out_path  = out_dir / f"selection_{type_tiers}_{date_str}.xlsx"
        sel.to_excel(out_path, index=False, engine="openpyxl")

        _format_excel_selection(out_path)

        return out_path

    def _run_lettres(self, eng, inputs, config, output_dir, cb):
        if cb: cb(15, "Chargement des données…")
        df = pd.read_excel(inputs["fichier_selection"])

        # Déterminer automatiquement le type de tiers depuis le nom du fichier
        filename = Path(inputs["fichier_selection"]).name.lower()
        is_client = "client" in filename or "clt" in filename
        canvas_name = "lc client.docx" if is_client else "lc fournisseur.docx"
        canvas_path = Path(__file__).parent / "canvas" / canvas_name

        if not canvas_path.exists():
            raise FileNotFoundError(f"Canvas introuvable : {canvas_path}")

        if cb: cb(30, f"Utilisation du canvas : {canvas_name}")
        template = eng.analyser_template(str(canvas_path))

        date_lettre = inputs.get("date_lettre", datetime.now().strftime("%d/%m/%Y"))

        # Extraire le nom de société depuis le fichier de sélection (première ligne ou config)
        nom_societe = "SOCIÉTÉ À AUDITER"  # Valeur par défaut
        try:
            # Essayer d'extraire depuis les données
            if "nom_societe" in df.columns and not df.empty:
                nom_societe = str(df["nom_societe"].iloc[0]) if pd.notna(df["nom_societe"].iloc[0]) else nom_societe
        except:
            pass

        config_gen = {**config.get("cabinet", {}),
                      "nom_societe": nom_societe, "date_cloture": date_lettre}

        col_nom   = _find_col(df, ["nom", "name", "nom_tiers", "libellé", "raison sociale"])
        col_code  = _find_col(df, ["code_tiers", "code", "compte", "n°"])
        col_solde = _find_col(df, ["solde", "montant", "closing", "credit", "balance"])

        if cb: cb(40, f"Génération de {len(df)} lettres…")
        out_dir = Path(output_dir)
        fichiers_gen = []

        for _, row in df.iterrows():
            tiers = {
                "nom":          str(row.get(col_nom, row.get(col_code, "Inconnu"))),
                "code_tiers":   str(row.get(col_code, "")),
                "solde_closing": float(row[col_solde]) if col_solde and pd.notna(row.get(col_solde)) else 0,
                "devise":       str(row.get("devise", "MAD")),
            }
            doc      = eng.remplir_template(template["chemin"], template["zones"], tiers, config_gen, date_lettre)
            nom_safe = eng.nettoyer_nom(tiers["nom"])
            nom_fic  = f"{nom_safe}_{tiers['code_tiers']}.docx"
            chemin_out = out_dir / nom_fic
            doc.save(str(chemin_out))
            fichiers_gen.append({"chemin": str(chemin_out), "nom": tiers["nom"], "code": tiers["code_tiers"]})

        if cb: cb(80, "Consolidation…")

        # Document consolidé
        if fichiers_gen:
            from docx import Document as DocxDoc
            nom_soc_safe = eng.nettoyer_nom(nom_societe) if nom_societe else "circularisation"
            type_tag = "Clients" if is_client else "Frs"
            nom_cons   = f"Circularisation_{type_tag}_{nom_soc_safe}.docx"
            doc_final  = DocxDoc(fichiers_gen[0]["chemin"])
            for fg in fichiers_gen[1:]:
                doc_final.add_page_break()
                doc_src = DocxDoc(fg["chemin"])
                for para in doc_src.paragraphs:
                    np_ = doc_final.add_paragraph()
                    np_.paragraph_format.alignment = para.paragraph_format.alignment
                    for run in para.runs:
                        nr = np_.add_run(run.text)
                        nr.bold, nr.italic, nr.underline = run.bold, run.italic, run.underline
                        if run.font.size: nr.font.size = run.font.size
                        if run.font.name: nr.font.name = run.font.name
            chemin_cons = out_dir / nom_cons
            doc_final.save(str(chemin_cons))

            import json
            etat = {"mapping_pages": [{"tiers": f["nom"], "code": f["code"]} for f in fichiers_gen],
                    "type": type_tag, "date": datetime.now().isoformat()}
            with open(out_dir / "etat_generation.json", "w", encoding="utf-8") as f:
                json.dump(etat, f, indent=2, ensure_ascii=False)

            return chemin_cons
        return None

    def _run_split(self, eng, inputs, config, output_dir, cb):
        try:
            from pypdf import PdfReader, PdfWriter
        except ImportError:
            raise ImportError(
                "Le module 'pypdf' est requis pour le split PDF.\n"
                "Installez-le : pip install pypdf"
            )
        import json

        if cb: cb(10, "Lecture du PDF…")
        pdf_path       = inputs.get("pdf_signe", "")
        pages_p_lettre = max(1, int(inputs.get("pages_par_lettre", 1) or 1))

        if not pdf_path or not Path(str(pdf_path)).exists():
            raise FileNotFoundError(f"PDF introuvable : {pdf_path}")

        reader      = PdfReader(str(pdf_path))
        total_pages = len(reader.pages)

        if total_pages == 0:
            raise ValueError("Le PDF est vide (0 pages).")

        if cb: cb(20, f"PDF chargé : {total_pages} page(s)…")

        # ── Mapping des tiers ────────────────────────────────
        mapping = []
        sel_path = inputs.get("fichier_selection", "")
        if sel_path and str(sel_path).strip() and Path(str(sel_path)).exists():
            df_sel   = pd.read_excel(sel_path)
            col_nom  = _find_col(df_sel, ["nom", "name", "nom_tiers", "libellé", "raison sociale", "intitulé", "libelle", "societe", "entreprise", "fournisseur", "client", "tiers"])
            col_code = _find_col(df_sel, ["code_tiers", "code", "compte", "n°", "id", "numero", "num", "reference", "ref"])
            for _, row in df_sel.iterrows():
                nom_val  = str(row[col_nom]).strip()  if col_nom  and pd.notna(row.get(col_nom))  else ""
                code_val = str(row[col_code]).strip() if col_code and pd.notna(row.get(col_code)) else ""
                if nom_val or code_val:
                    mapping.append({"tiers": nom_val or code_val, "code": code_val})

        # Mapping automatique si aucune sélection fournie
        if not mapping:
            nb_lettres = max(1, total_pages // pages_p_lettre)
            mapping = [{"tiers": f"Lettre_{i+1}", "code": f"{i+1:03d}"}
                       for i in range(nb_lettres)]

        if cb: cb(30, f"Découpage : {total_pages} page(s) → {len(mapping)} lettre(s)…")

        # ── Sous-dossier dédié avec timestamp ─────────────────
        ts      = datetime.now().strftime("%Y%m%d_%H%M%S")
        out_dir = Path(output_dir) / f"split_{ts}"
        out_dir.mkdir(parents=True, exist_ok=True)

        page_courante  = 0
        fichiers_split = []
        noms_utilises  = set()  # éviter les doublons de noms de fichiers

        for i, tiers_info in enumerate(mapping):
            nom  = tiers_info.get("tiers", f"Lettre_{i+1}") or f"Lettre_{i+1}"
            code = tiers_info.get("code", f"{i+1:03d}") or f"{i+1:03d}"

            debut = page_courante
            fin   = min(debut + pages_p_lettre, total_pages)
            if debut >= total_pages:
                break

            writer = PdfWriter()
            for p in range(debut, fin):
                writer.add_page(reader.pages[p])

            nom_safe = eng.nettoyer_nom(nom)
            if not nom_safe:
                nom_safe = f"tiers_{i+1}"

            # Éviter les collisions de noms
            if nom_safe and nom_safe != code:
                base = nom_safe
            else:
                base = f"{code}_{nom_safe}" if code else nom_safe
            
            nom_fic = f"{base}.pdf"
            if nom_fic in noms_utilises:
                nom_fic = f"{base}_{i+1}.pdf"
            noms_utilises.add(nom_fic)

            chemin_out = out_dir / nom_fic
            with open(chemin_out, "wb") as f:
                writer.write(f)

            taille_kb = chemin_out.stat().st_size // 1024
            fichiers_split.append({
                "chemin": str(chemin_out),
                "tiers":  nom,
                "code":   code,
                "pages":  f"{debut+1}-{fin}",
                "taille_kb": taille_kb,
            })
            page_courante = fin

            if cb:
                pct = 30 + int(65 * (i + 1) / len(mapping))
                cb(pct, f"{i+1}/{len(mapping)} — {nom[:35]}")

        # Pages restantes non couvertes
        restantes = total_pages - page_courante
        if restantes > 0:
            writer = PdfWriter()
            for p in range(page_courante, total_pages):
                writer.add_page(reader.pages[p])
            nom_reste = f"Pages_restantes_{page_courante+1}-{total_pages}.pdf"
            with open(out_dir / nom_reste, "wb") as f:
                writer.write(f)
            fichiers_split.append({
                "chemin": str(out_dir / nom_reste),
                "tiers": "Pages restantes",
                "code": "RESTE",
                "pages": f"{page_courante+1}-{total_pages}",
            })

        # Sauvegarder le manifeste JSON dans le sous-dossier
        with open(out_dir / "split_result.json", "w", encoding="utf-8") as f:
            json.dump(fichiers_split, f, indent=2, ensure_ascii=False)

        if cb: cb(98, f"{len(fichiers_split)} PDF créés dans {out_dir.name}")

        # ── Retourner le DOSSIER contenant les PDFs ───────────
        # (pas le JSON — l'UI ouvrira le dossier dans l'explorateur)
        return out_dir

    def _run_emails(self, eng, inputs, config, output_dir, cb):
        import json
        from email.mime.multipart import MIMEMultipart
        from email.mime.text import MIMEText
        from email.mime.application import MIMEApplication
        from email.utils import formatdate

        if cb: cb(10, "Chargement des données…")
        df         = pd.read_excel(inputs["fichier_selection"])
        pdf_dir    = Path(inputs["dossier_pdf"])
        nom_societe = inputs.get("nom_societe", "")
        date_arrete = inputs.get("date_arrete", "31/12/2024")
        type_tiers  = inputs.get("type_tiers", "Clients")
        nom_cabinet = config.get("cabinet", {}).get("nom", "AuditPro Cabinet")
        email_exp   = config.get("cabinet", {}).get("email_expediteur", "audit@auditpro.local")

        col_nom   = _find_col(df, ["nom", "name", "nom_tiers", "libellé"])
        col_code  = _find_col(df, ["code_tiers", "code", "compte"])
        col_email = _find_col(df, ["email", "mail", "e-mail", "courriel"])
        col_solde = _find_col(df, ["solde", "montant", "closing", "balance"])

        SUJET = f"Circularisation – Confirmation de solde – {nom_societe}"
        CORPS = (
            f"Bonjour,\n\n"
            f"Dans le cadre de notre mission de commissariat aux comptes pour la société "
            f"« {nom_societe} » au titre de l'exercice clos le {date_arrete}, nous vous "
            f"prions de bien vouloir répondre à la lettre de circularisation ci-jointe.\n\n"
            f"En vous remerciant par avance pour votre retour,\n\n"
            f"Bien cordialement,\n{nom_cabinet}"
        )

        out_dir = Path(output_dir)
        generes, manquants = 0, []

        if cb: cb(20, f"Génération de {len(df)} emails…")

        for i, row in df.iterrows():
            nom   = str(row.get(col_nom, row.get(col_code, f"Tiers_{i}")))
            code  = str(row.get(col_code, ""))
            email = str(row.get(col_email, "")) if col_email else ""
            montant = float(row[col_solde]) if col_solde and pd.notna(row.get(col_solde)) else 0

            nom_safe = eng.nettoyer_nom(nom)
            pdf_candidates = (
                list(pdf_dir.glob(f"*{nom_safe}*.pdf")) +
                list(pdf_dir.glob(f"*{code}*.pdf"))
            )
            pdf_path = pdf_candidates[0] if pdf_candidates else None

            if not pdf_path:
                manquants.append({"tiers": nom, "motif": "PDF manquant"})
                continue
            if not email or "@" not in email:
                manquants.append({"tiers": nom, "motif": "Email manquant"})
                email = ""

            msg = MIMEMultipart()
            msg["Subject"] = SUJET
            msg["To"]      = email
            msg["From"]    = email_exp
            msg["Date"]    = formatdate(localtime=True)
            msg.attach(MIMEText(CORPS, "plain", "utf-8"))

            with open(pdf_path, "rb") as f:
                pj = MIMEApplication(f.read(), _subtype="pdf")
                pj.add_header("Content-Disposition", "attachment", filename=pdf_path.name)
                msg.attach(pj)

            eml_name = f"Email_{code}_{nom_safe}.eml"
            with open(out_dir / eml_name, "w", encoding="utf-8") as f:
                f.write(msg.as_string())
            generes += 1

            if cb: cb(20 + int(70 * (i + 1) / len(df)), f"{i+1}/{len(df)} — {nom[:30]}")

        if manquants:
            rapport = out_dir / "emails_manquants.xlsx"
            pd.DataFrame(manquants).to_excel(rapport, index=False)

        rapport_final = out_dir / f"rapport_emails_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
        pd.DataFrame({
            "Statut":  ["Générés", "Avec problème"],
            "Nombre":  [generes, len(manquants)],
        }).to_excel(rapport_final, index=False)

        return rapport_final


# ── Helpers internes ───────────────────────────────────────────

def _find_col(df, keywords):
    """Trouve la première colonne correspondant à un mot-clé (insensible casse)."""
    for kw in keywords:
        for col in df.columns:
            if kw.lower() in str(col).lower():
                return col
    return None


def _auto_mapper(df, type_tiers, eng, config):
    """Mapping automatique sans interaction (utilise les keywords du script)."""
    rules = config.get("mapping", {}).get(type_tiers, {})
    mapping = {}
    for role, keywords in rules.items():
        col = eng.trouver_colonne(df, keywords)
        if col:
            mapping[role] = col
    return mapping


def _format_excel_selection(path: Path):
    """Mise en forme Excel de la sélection."""
    try:
        from openpyxl import load_workbook
        from openpyxl.styles import Font, PatternFill, Border, Side
        wb = load_workbook(path)
        ws = wb.active
        hdr_fill = PatternFill("solid", fgColor="4B286D")   # violet GT
        hdr_font = Font(bold=True, color="FFFFFF", name="Arial", size=10)
        vert_fill = PatternFill("solid", fgColor="EDE7F6")
        bd = Border(left=Side("thin"), right=Side("thin"),
                    top=Side("thin"), bottom=Side("thin"))
        for cell in ws[1]:
            cell.fill, cell.font, cell.border = hdr_fill, hdr_font, bd
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row, max_col=ws.max_column):
            for cell in row:
                cell.border = bd
                cell.font   = Font(name="Arial", size=10)
                hdr = str(ws.cell(1, cell.column).value or "").lower()
                if hdr in ("email", "adresse") and not cell.value:
                    cell.fill = vert_fill
        ws.freeze_panes = "A2"
        for col in ws.columns:
            ws.column_dimensions[col[0].column_letter].width = max(
                12, min(42, max(len(str(c.value or "")) for c in col) + 2)
            )
        wb.save(path)
    except Exception:
        pass  # Mise en forme optionnelle
