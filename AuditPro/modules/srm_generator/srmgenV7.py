#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
SRMGEN - Generateur automatique de SRM
Tableaux Word natifs + commentaires varies + guide remplacement images
Usage: python srmgen.py | python srmgen.py fichier.xlsx | python srmgen.py -i
"""
import os, sys, random, argparse
from datetime import datetime
from pathlib import Path
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE_DIR = Path(__file__).parent
INPUT_DIR = BASE_DIR / "input"
OUTPUT_DIR = BASE_DIR / "output"

# Synonymes
S_AUG = ["une augmentation", "une hausse", "une progression", "un accroissement"]
S_DIM = ["une diminution", "une baisse", "un repli", "un recul"]
P_DEB = [
    "Le poste {L} pr\u00e9sente un solde de",
    "Concernant le poste {L}, il pr\u00e9sente",
    "Le compte {L} fait ressortir un montant global de",
    "Les compte des {L} affiche un solde de",
    "Le compte des {L} laisse appara\u00eetre un solde de",
    "Le compte {L} montre un solde de",
    "S'agissant du poste {L}, il pr\u00e9sente un solde global de",
    "Relatif au poste {L}, il pr\u00e9sente un montant total de",
]
P_RES = [
    "{L} s'\u00e9tablit \u00e0",
    "{L} totalise un montant de",
    "{L} s'\u00e9l\u00e8ve \u00e0",
    "{L} a enregistr\u00e9 un solde de",
    "{L} ressort \u00e0",
]
P_INTRO = [
    "Le d\u00e9tail se pr\u00e9sente comme suit :",
    "Cette variation se d\u00e9taille comme suit :",
    "Par ailleurs, le d\u00e9tail se pr\u00e9sente comme suit :",
    "Les \u00e9l\u00e9ments suivants permettent d'expliquer cette variation :",
]
_ud, _ur = [], []
def _pick(pool, used):
    a = [p for p in pool if p not in used]
    if not a: used.clear(); a = pool
    c = random.choice(a); used.append(c); return c

def fmt(n):
    if n is None: return "\u2013"
    try:
        v = round(float(n))
        return ("-" if v < 0 else "") + f"{abs(v):,}".replace(",", " ")
    except: return str(n)

def sf(v, d=None):
    if v is None: return d
    if isinstance(v, datetime): return d
    if isinstance(v, str) and v.strip() in ("", "-", "\u2013", "nan", "None", "#REF!"): return d
    try: return float(v)
    except: return d

def log(m, l="INFO"):
    icons = {"INFO": "\u2139\ufe0f", "OK": "\u2705", "WARN": "\u26a0\ufe0f", "ERR": "\u274c", "TAB": "\U0001f4ca"}
    print(f"  {icons.get(l, chr(8226))} {m}")

# Couleurs theme Excel
TH = {0:"FFFFFF",1:"000000",2:"44546A",3:"E7E6E6",4:"4472C4",5:"ED7D31",6:"A5A5A5",7:"FFC000",8:"5B9BD5",9:"70AD47"}

def c2h(co, d=None):
    if co is None: return d
    try:
        if co.type == "rgb" and co.rgb:
            rgb = str(co.rgb)
            if len(rgb) == 8:
                if rgb[:2] == "00": return d
                return rgb[2:]
            if len(rgb) == 6: return rgb
        if co.type == "theme": return TH.get(co.theme, d)
        if co.type == "indexed":
            m = {0:"000000",1:"FFFFFF",2:"FF0000",3:"00FF00",4:"0000FF",5:"FFFF00",8:"000000",9:"FFFFFF",64:"000000"}
            return m.get(co.indexed, d)
    except: pass
    return d

def h2rgb(h):
    if not h or len(h) != 6: return RGBColor(0, 0, 0)
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def gcf(cell):
    f = cell.font; fill = cell.fill
    bold = bool(f.bold) if f else False
    size = float(f.size) if f and f.size else 9.0
    fn = f.name if f and f.name else "Arial"
    fc = c2h(f.color, "000000") if f else "000000"
    bg = None
    if fill and fill.fill_type and fill.fill_type != "none":
        bg = c2h(fill.fgColor)
        if bg == "000000": bg = None
    al = cell.alignment.horizontal if cell.alignment else None
    nf = cell.number_format or "General"
    return {"bold":bold,"size":size,"font":fn,"color":fc,"bg":bg,"align":al,"nfmt":nf}

def fval(value, nfmt):
    if value is None: return ""
    if isinstance(value, datetime): return value.strftime("%d/%m/%Y")
    if isinstance(value, str): return value.strip()
    try: fv = float(value)
    except: return str(value)
    if "%" in nfmt:
        pct = fv * 100 if abs(fv) <= 2.0 else fv
        sign = "-" if pct < 0 else ""
        dec = 1 if "0.0" in nfmt else 0
        return f"{sign}{abs(pct):.{dec}f} %".replace(".", ",")
    if "#,##0" in nfmt or "# ##0" in nfmt:
        v = round(fv); sign = "-" if v < 0 else ""
        return sign + f"{abs(v):,}".replace(",", " ")
    if fv == int(fv): return str(int(fv))
    return f"{fv:.2f}".replace(".", ",")

def is_mk(v):
    if v is None: return False
    s = str(v).strip().lower()
    return "milliers" in s or "en kmad" in s

def detect_blocks(ws):
    mk = []; mr = ws.max_row or 1; mc = ws.max_column or 1
    # Détecter les cellules avec "milliers" ou "kmad"
    for row in ws.iter_rows(min_row=1, max_row=mr):
        for cell in row:
            if cell.value and is_mk(cell.value): mk.append((cell.row, cell.column))

    # Si aucun marqueur trouvé, essayer de détecter automatiquement les blocs de données
    if not mk:
        log("Aucun marqueur 'milliers' trouvé, tentative de détection automatique...", "WARN")
        for r in range(1, min(100, mr)):
            for c in range(1, min(20, mc)):
                cell = ws.cell(r, c)
                # Chercher les headers textuels suivis de données numériques
                if cell.value and isinstance(cell.value, str) and len(cell.value.strip()) > 3:
                    # Vérifier s'il y a des données numériques dans les lignes suivantes
                    has_numeric = False
                    for rr in range(r+1, min(r+10, mr)):
                        for cc in range(c, min(c+10, mc)):
                            val = ws.cell(rr, cc).value
                            if isinstance(val, (int, float)) and abs(val) > 10:  # Valeur significative
                                has_numeric = True
                                break
                        if has_numeric: break
                    if has_numeric:
                        mk.append((r, c))
                        continue

                # Chercher aussi les blocs avec des valeurs numériques élevées
                if cell.value and isinstance(cell.value, (int, float)) and abs(cell.value) > 100:
                    # Chercher si cette ligne contient des données numériques
                    num_count = 0
                    for cc in range(c, min(c+10, mc)):
                        val = ws.cell(r, cc).value
                        if isinstance(val, (int, float)): num_count += 1
                    if num_count >= 2:  # Au moins 2 colonnes numériques
                        mk.append((r, c))

    blocks = []
    for (r0, c0) in mk:
        ec = c0
        for c in range(c0+1, min(c0+25, mc+2)):
            has = any(ws.cell(r,c).value is not None and str(ws.cell(r,c).value).strip() not in ("","#REF!") for r in range(r0, min(r0+40, mr+1)))
            if has: ec = c
            else:
                lt = get_column_letter(c); cd = ws.column_dimensions.get(lt)
                if cd and cd.width is not None and cd.width < 1.5: continue
                break
        er = r0; emp = 0
        for r in range(r0+1, min(r0+80, mr+1)):
            has = any(ws.cell(r,c).value is not None and str(ws.cell(r,c).value).strip() not in ("","#REF!") for c in range(c0, ec+1))
            if has: er = r; emp = 0
            else:
                emp += 1
                if emp >= 2: break
        if er > r0 + 1: blocks.append({"r0":r0,"c0":c0,"r2":er,"c2":ec})
    u = []
    for b in blocks:
        if not any(abs(b["r0"]-x["r0"])<3 and abs(b["c0"]-x["c0"])<3 for x in u): u.append(b)
    return sorted(u, key=lambda x: (x["r0"], x["c0"]))

CPC_KW = ["produits d'exploitation","charges d'exploitation","r\u00e9sultat d'exploitation","r\u00e9sultat financier","r\u00e9sultat non courant","r\u00e9sultat net","r\u00e9sultat avant","imp\u00f4t sur le r\u00e9sultat","ventes de biens","produits financiers","charges financi\u00e8res"]
BILAN_KW = ["total de l'actif","total du passif","actif immobilis\u00e9","actif circulant","passif circulant","financement permanent","tr\u00e9sorerie","capitaux propres","immobilisation"]
DETAIL_KW = ["clients","fournisseurs","stocks","immobilisations","amortissements","provisions","emprunts","dettes","creances","charges","produits","ventes","achats","salaires","impots","tva","cnss"]

def classify(ws, b):
    t = []
    for r in range(b["r0"], b["r2"]+1):
        for c in range(b["c0"], b["c2"]+1):
            v = ws.cell(r,c).value
            if isinstance(v, str): t.append(v.lower())
    comb = " ".join(t)

    # Comptage strict des mots-clés complets (pas partiels)
    cs = sum(1 for k in CPC_KW if k in comb)
    bs = sum(1 for k in BILAN_KW if k in comb)
    ds = sum(1 for k in DETAIL_KW if k in comb)

    # Logique de classification hiérarchique
    # 1. Si contient des termes principaux de bilan, c'est bilan
    if bs >= 1: return "bilan"
    # 2. Si contient des termes principaux de CPC, c'est CPC
    if cs >= 1: return "cpc"
    # 3. Si contient des termes de détail OU est petit (peu de lignes), c'est détail
    if ds >= 1 or (b["r2"] - b["r0"]) <= 5: return "detail"
    # 4. Par défaut, détail
    return "detail"

def extract_total(ws, b):
    dates = []
    # Chercher les dates dans tout le bloc, pas seulement les 3 premières lignes
    for r in range(b["r0"], b["r2"]+1):
        for c in range(b["c0"], b["c2"]+1):
            v = ws.cell(r,c).value
            if isinstance(v, datetime): dates.append((c, v))
    # Trier par colonne pour avoir N puis N-1
    dates.sort(key=lambda x: x[0])

    dn = dates[0][1] if len(dates) >= 1 else datetime(datetime.today().year, 12, 31)
    dn1 = dates[1][1] if len(dates) >= 2 else datetime(datetime.today().year-1, 12, 31)
    cn = dates[0][0] if len(dates) >= 1 else None
    cn1 = dates[1][0] if len(dates) >= 2 else None
    best = None
    for r in range(b["r2"], b["r0"], -1):
        for c in range(b["c0"], b["c2"]+1):
            cell = ws.cell(r,c); v = cell.value
            if isinstance(v, str) and len(v.strip()) > 2:
                ib = cell.font and cell.font.bold
                ik = any(k in v.lower() for k in ["total","r\u00e9sultat","charges d'","produits d'"])
                if ib or ik:
                    nv = sf(ws.cell(r,cn).value) if cn else None
                    n1v = sf(ws.cell(r,cn1).value) if cn1 else None
                    var = (nv-n1v) if nv is not None and n1v is not None else None
                    best = {"label":v.strip(),"n":nv,"n1":n1v,"var":var,"date_n":dn,"date_n1":dn1}
                    break
        if best: break

    # Si aucun total trouvé avec les critères habituels, essayer de trouver une ligne avec des valeurs numériques significatives
    if best is None:
        # Pour les sous-comptes, chercher la ligne avec la plus grande valeur absolue
        max_val = 0
        max_row = None
        max_col = None
        for r in range(b["r0"], b["r2"]+1):
            for c in range(b["c0"], b["c2"]+1):
                cell = ws.cell(r,c); v = cell.value
                if isinstance(v, (int, float)) and abs(v) > max_val:
                    max_val = abs(v)
                    max_row = r
                    max_col = c

        if max_row is not None and max_val > 10:  # Valeur significative
            nv = sf(ws.cell(max_row, cn).value) if cn else sf(max_val)
            n1v = sf(ws.cell(max_row, cn1).value) if cn1 else None
            var = (nv-n1v) if nv is not None and n1v is not None else None

            # Chercher le label dans la même ligne ou ligne au-dessus
            label = None
            for c in range(b["c0"], max_col):  # Chercher à gauche de la valeur
                lv = ws.cell(max_row, c).value
                if isinstance(lv, str) and len(lv.strip()) > 2:
                    label = lv.strip()
                    break
            if not label:  # Chercher dans la ligne au-dessus
                for c in range(b["c0"], b["c2"]+1):
                    lv = ws.cell(max_row-1, c).value
                    if isinstance(lv, str) and len(lv.strip()) > 2:
                        label = lv.strip()
                        break

            best = {"label":label or f"Ligne {max_row-b['r0']+1}","n":nv,"n1":n1v,"var":var,"date_n":dn,"date_n1":dn1}

    return best

def gen_com(label, n, n1, var, dn, dn1):
    ds = dn.strftime("%d/%m/%Y") if hasattr(dn,"strftime") else str(dn)
    d1 = dn1.strftime("%d/%m/%Y") if hasattr(dn1,"strftime") else str(dn1)
    nf, n1f = sf(n,0), sf(n1,0); vf = sf(var, nf-n1f)
    ir = "r\u00e9sultat" in label.lower()
    deb = _pick(P_RES if ir else P_DEB, _ur if ir else _ud).format(L=label)
    base = f"{deb} {fmt(nf)} KMAD au {ds}, contre {fmt(n1f)} KMAD au {d1}"
    if abs(n1f) < 0.01: return f"{base}."
    pct = abs(round(((nf-n1f)/abs(n1f))*100, 1))
    if abs(vf) < 1: return f"{base}. Ce poste est rest\u00e9 globalement stable sur la p\u00e9riode."
    syn = random.choice(S_AUG if vf > 0 else S_DIM)
    txt = f"{base}, soit {syn} de {fmt(abs(round(vf)))} KMAD ({pct} %)."
    return txt

def set_shd(cell, hx):
    tcPr = cell._tc.get_or_add_tcPr()
    old = tcPr.find(qn('w:shd'))
    if old is not None: tcPr.remove(old)
    el = tcPr.makeelement(qn('w:shd'), {qn('w:val'):'clear',qn('w:color'):'auto',qn('w:fill'):hx})
    tcPr.append(el)

def add_table(doc, ws, b, tnum):
    r0, c0, r2, c2 = b["r0"], b["c0"], b["r2"], b["c2"]
    ac = [c for c in range(c0, c2+1) if any(ws.cell(r,c).value is not None and str(ws.cell(r,c).value).strip() not in ("","#REF!") for r in range(r0, r2+1))]
    if not ac: return
    ar = [r for r in range(r0, r2+1) if any(ws.cell(r,c).value is not None and str(ws.cell(r,c).value).strip() not in ("","#REF!") for c in ac)]
    if not ar: return
    p = doc.add_paragraph()
    run = p.add_run(f"[TABLE_{tnum:03d}]")
    run.font.size = Pt(6); run.font.color.rgb = RGBColor(0xCC,0xCC,0xCC)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table = doc.add_table(rows=len(ar), cols=len(ac))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for ri, row_idx in enumerate(ar):
        for ci, col_idx in enumerate(ac):
            cx = ws.cell(row_idx, col_idx)
            cw = table.rows[ri].cells[ci]
            fi = gcf(cx)
            vt = fval(cx.value, fi["nfmt"])
            cw.text = ""
            p = cw.paragraphs[0]
            run = p.add_run(vt)
            run.font.name = fi["font"]; run.font.size = Pt(fi["size"])
            run.bold = fi["bold"]; run.font.color.rgb = h2rgb(fi["color"])
            am = {"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER,"centerContinuous":WD_ALIGN_PARAGRAPH.CENTER,"right":WD_ALIGN_PARAGRAPH.RIGHT}
            if fi["align"] and fi["align"] in am: p.alignment = am[fi["align"]]
            elif isinstance(cx.value, (int, float)): p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else: p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if fi["bg"]: set_shd(cw, fi["bg"])
            p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)

def build_srm(blocks, out_path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2); sec.right_margin = Cm(2)
    s = doc.styles["Normal"]; s.font.name = "Arial"; s.font.size = Pt(9)
    p = doc.add_paragraph()
    run = p.add_run("E. Revue analytique des comptes")
    run.bold = True; run.font.size = Pt(14); run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0x1F,0x4E,0x79)
    doc.add_paragraph()
    cc = None; ai = 0; al = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"; tn = 0
    for b in blocks:
        ws = b.get("ws")  # worksheet propre à ce bloc (fix multi-feuilles)
        cat = b["category"]
        if cat != cc:
            cc = cat; ai = 0; p = doc.add_paragraph()
            lb = {"cpc":"Comptes de r\u00e9sultat","bilan":"Bilan","detail":"D\u00e9tails par rubrique"}[cat]
            run = p.add_run(lb)
            run.bold = True; run.font.size = Pt(12); run.font.name = "Arial"
            run.font.color.rgb = RGBColor(0x1F,0x4E,0x79)
            doc.add_paragraph()
        total = extract_total(ws, b); tn += 1
        if total and total["n"] is not None:
            lt = al[ai%26]; ai += 1
            p = doc.add_paragraph()
            run = p.add_run(f"{lt}. {total['label']}")
            run.bold = True; run.font.size = Pt(10); run.font.name = "Arial"
            run.font.color.rgb = RGBColor(0x1F,0x4E,0x79)
            com = gen_com(total["label"], total["n"], total["n1"], total["var"], total["date_n"], total["date_n1"])
            p = doc.add_paragraph()
            for part in com.split("\n"):
                run = p.add_run(part); run.font.size = Pt(9); run.font.name = "Arial"
                if part != com.split("\n")[-1]: run.add_break()
            p = doc.add_paragraph()
            run = p.add_run(random.choice(P_INTRO))
            run.font.size = Pt(9); run.font.name = "Arial"
        add_table(doc, ws, b, tn)
        doc.add_paragraph()
    # Guide
    doc.add_page_break()
    p = doc.add_paragraph()
    run = p.add_run("GUIDE : Remplacer les tableaux par des images Excel")
    run.bold = True; run.font.size = Pt(14); run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0xC0,0x00,0x00)
    guide = [
        "",
        "Ce document contient des tableaux Word natifs. Pour obtenir un rendu",
        "identique a Excel (pixel-perfect), remplacez-les par des images :",
        "",
        "METHODE (pour chaque tableau) :",
        "1. Dans Excel, selectionner le tableau (les cellules du bloc)",
        "2. Copier (Ctrl+C)",
        "3. Dans ce Word, selectionner le tableau (clic sur la croix en haut a gauche)",
        "4. Supprimer le tableau (touche Suppr)",
        "5. Collage special : Ctrl+Alt+V puis Image (Metafichier ameliore)",
        "",
        "ASTUCE : Chaque tableau est precede d un marqueur gris [TABLE_001], [TABLE_002]...",
        "Utilisez Ctrl+H pour les retrouver. Supprimez le marqueur apres remplacement.",
        "",
        f"Nombre total de tableaux : {tn}",
    ]
    for line in guide:
        p = doc.add_paragraph()
        if line.startswith("METHODE") or line.startswith("ASTUCE") or line.startswith("Nombre"):
            run = p.add_run(line); run.bold = True; run.font.size = Pt(10); run.font.name = "Arial"
        else:
            run = p.add_run(line); run.font.size = Pt(9); run.font.name = "Arial"

    # Sauvegarde du document
    try:
        log(f"Sauvegarde du document : {out_path}")
        doc.save(str(out_path))
        log("Document sauvegardé avec succès")
        return out_path
    except Exception as e:
        log(f"ERREUR lors de la sauvegarde : {e}", "ERR")
        raise

def process_file(fp, od=None):
    fp = Path(fp); od = Path(od) if od else OUTPUT_DIR; od.mkdir(parents=True, exist_ok=True)
    print(f"\n{'='*60}"); print(f"  SRMGEN \u2014 Generation du Summary Review Memorandum"); print(f"{'='*60}")
    log(f"Fichier : {fp.name}")
    log(f"Répertoire de sortie : {od}")

    # Vérifier que le fichier existe
    if not fp.exists():
        log(f"ERREUR : Fichier introuvable : {fp}", "ERR")
        return None

    # Vérifier les permissions du répertoire de sortie
    if not od.exists():
        log(f"Création du répertoire : {od}", "INFO")
        od.mkdir(parents=True, exist_ok=True)

    try:
        test_file = od / ".srm_test"
        with open(test_file, 'w') as f:
            f.write("test")
        test_file.unlink()
        log("Permissions d'écriture : OK", "OK")
    except Exception as e:
        log(f"Permissions d'écriture : ÉCHEC - {e}", "ERR")
        return None

    wb = load_workbook(str(fp), data_only=True); log(f"Feuilles : {wb.sheetnames}")
    ab = []
    for sn in wb.sheetnames:
        ws = wb[sn]; bl = detect_blocks(ws)
        if bl:
            log(f"Feuille '{sn}' : {len(bl)} tableaux", "TAB")
            for b in bl:
                b["ws"] = ws  # stocker le worksheet dans le bloc (fix multi-feuilles)
                b["category"] = classify(ws, b); t = extract_total(ws, b)
                ti = t["label"][:50] if t else f"R{b['r0']}:C{b['c0']}"
                log(f"  \u2192 [{b['category'].upper():6s}] {ti}")
            ab.extend(bl)
    if not ab: log("Aucun tableau detecte.", "ERR"); return None
    o = {"cpc":0,"bilan":1,"detail":2}
    ab.sort(key=lambda b: (o[b["category"]], b["r0"], b["c0"]))
    cs = {}
    for b in ab: cs[b["category"]] = cs.get(b["category"], 0) + 1
    log(f"Total : {len(ab)} tableaux \u2014 " + ", ".join(f"{k.upper()}: {v}" for k,v in cs.items()), "OK")
    stem = fp.stem
    for px in ("tab_srm_","tableau_srm_","tab_","tableau_"):
        if stem.lower().startswith(px): stem = stem[len(px):]; break
    stem = "SRM_" + stem
    t0 = extract_total(ab[0]["ws"], ab[0])  # utiliser le ws propre au premier bloc
    yr = t0["date_n"].year if t0 and hasattr(t0["date_n"],"year") else datetime.today().year
    op = od / f"{stem}_{yr}.docx"
    log(f"Fichier de sortie prévu : {op}")
    log("Generation Word (tableaux natifs + commentaires)...")

    try:
        result = build_srm(ab, op)  # plus de ws global, chaque bloc a son propre ws
        log(f"Document Word créé : {result}")
        log(f"SRM genere : {op}", "OK")
        log(f"Pour remplacer par images Excel, voir le guide en fin de document.")
        return op
    except Exception as e:
        log(f"ERREUR lors de la génération Word : {e}", "ERR")
        import traceback
        log(f"Traceback : {traceback.format_exc()}", "ERR")
        raise

def interactive_mode():
    print(f"\n{'='*60}"); print(f"  SRMGEN \u2014 Mode Interactif"); print(f"{'='*60}\n")
    INPUT_DIR.mkdir(parents=True, exist_ok=True); files = sorted(INPUT_DIR.glob("*.xlsx"))
    if not files: log(f"Aucun .xlsx dans : {INPUT_DIR}", "ERR"); return
    print("  Fichiers disponibles :\n")
    for i, f in enumerate(files): print(f"    [{i+1}] {f.name}  ({f.stat().st_size/1024:.0f} KB)")
    print(f"\n    [0] Quitter\n")
    while True:
        try:
            ch = input("  Votre choix : ").strip()
            if ch == "0": print("  Au revoir !"); return
            idx = int(ch) - 1
            if 0 <= idx < len(files): break
        except: pass
        print(f"  Choix invalide (1-{len(files)} ou 0)")
    sel = files[idx]; print(f"\n  Previsualisation : {sel.name}\n")
    wb = load_workbook(str(sel), data_only=True)
    for sn in wb.sheetnames:
        ws = wb[sn]; bl = detect_blocks(ws)
        if bl:
            for b_ in bl:
                b_["category"] = classify(ws, b_); t = extract_total(ws, b_)
                ti = t["label"][:45] if t else "\u2014"
                ns = fmt(t["n"]) if t and t["n"] is not None else "\u2014"
                ic = {"cpc":"\U0001f4ca","bilan":"\U0001f4cb","detail":"\U0001f4dd"}[b_["category"]]
                print(f"    {ic} [{b_['category'].upper():6s}] {ti:45s} N={ns}")
    wb.close(); print()
    cn = input("  Generer le SRM ? (O/n) : ").strip().lower()
    if cn in ("n","non"): print("  Annule."); return
    process_file(sel)

def main():
    pa = argparse.ArgumentParser(description="SRMGEN")
    pa.add_argument("file", nargs="?"); pa.add_argument("--interactive", "-i", action="store_true")
    args = pa.parse_args()
    if args.interactive: interactive_mode(); return
    if args.file: process_file(args.file); return
    INPUT_DIR.mkdir(parents=True, exist_ok=True); files = sorted(INPUT_DIR.glob("*.xlsx"))
    if not files: log(f"Aucun .xlsx dans : {INPUT_DIR}", "ERR"); log("Usage : python srmgen.py [fichier.xlsx] ou python srmgen.py -i"); sys.exit(1)
    if len(files) == 1: process_file(files[0])
    else: log(f"{len(files)} fichiers. Utilisez -i pour choisir."); process_file(files[0])

if __name__ == "__main__": main()