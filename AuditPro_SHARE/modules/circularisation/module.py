"""
Module : Circularisation des Tiers (NEP 505 / ISA 505)
Workflow en 4 étapes indépendantes — chaque étape a ses propres inputs.
"""
import os
import importlib.util
import io
import re
import sys
from contextlib import redirect_stdout
from pathlib import Path
from datetime import datetime

import pandas as pd

from modules.base_module import BaseModule, ModuleInput, ModuleResult

# ── Définition des 4 étapes ────────────────────────────────────

ETAPES = [
    {
        "id":          "selection",
        "label":       "Étape 1 — Sélection des tiers",
        "description": "Charge une balance auxiliaire ou un GL et sélectionne les tiers à circulariser par cumul décroissant (85-95 %).\n\nSortie : fichier Excel des tiers sélectionnés (à compléter avec emails/adresses).",
        "icon":        "1",
    },
    {
        "id":          "lettres",
        "label":       "Étape 2 — Génération des lettres",
        "description": "Génère les lettres de circularisation Word individuelles à partir du fichier de sélection et d'un modèle Word (canvas).\n\nSortie : lettres .docx individuelles + document consolidé.",
        "icon":        "2",
    },
    {
        "id":          "split",
        "label":       "Étape 3 — Split PDF signé",
        "description": "Découpe le PDF signé (imprimé, signé, scanné puis remis en PDF) en autant de fichiers individuels que de tiers.\n\nSortie : PDF individuels nommés par tiers.",
        "icon":        "3",
    },
    {
        "id":          "emails",
        "label":       "Étape 4 — Génération des emails",
        "description": "Génère les fichiers .eml prêts à l'envoi, avec la lettre PDF en pièce jointe et le corps personnalisé par tiers.\n\nSortie : fichiers .eml + rapport des emails manquants.",
        "icon":        "4",
    },
]

# ── Inputs par étape ───────────────────────────────────────────

INPUTS_PAR_ETAPE = {

    "selection": [
        ModuleInput(
            key="balance_aux",
            label="Balance auxiliaire / Grand Livre",
            input_type="file",
            extensions=[".xlsx", ".xls", ".csv"],
            tooltip="Fichier Excel avec les soldes ou mouvements par tiers (clients ou fournisseurs)",
        ),
        ModuleInput(
            key="type_tiers",
            label="Type de tiers",
            input_type="combo",
            options=["Fournisseurs", "Clients"],
            default="Fournisseurs",
            required=True,
        ),
        ModuleInput(
            key="couverture_min",
            label="Couverture minimale (%)",
            input_type="number",
            default=85,
            required=False,
            tooltip="Ex : 85 → sélectionner au moins 85 % du montant total",
        ),
        ModuleInput(
            key="couverture_max",
            label="Couverture maximale (%)",
            input_type="number",
            default=95,
            required=False,
            tooltip="Ex : 95 → ne pas dépasser 95 % du montant total",
        ),
    ],

    "lettres": [
        ModuleInput(
            key="fichier_selection",
            label="Fichier de sélection (Excel)",
            input_type="file",
            extensions=[".xlsx", ".xls"],
            tooltip="Fichier produit par l'étape 1 (selection/output/)",
        ),
        ModuleInput(
            key="date_lettre",
            label="Date de la lettre",
            input_type="text",
            default=datetime.now().strftime("%d/%m/%Y"),
            tooltip="Date à afficher dans les lettres (JJ/MM/AAAA)",
        ),
    ],

    "split": [
        ModuleInput(
            key="pdf_signe",
            label="PDF signé à découper",
            input_type="file",
            extensions=[".pdf"],
            tooltip="PDF contenant toutes les lettres signées (une ou plusieurs pages par lettre)",
        ),
    ],

    "emails": [
        ModuleInput(
            key="fichier_selection",
            label="Fichier de sélection avec emails",
            input_type="file",
            extensions=[".xlsx", ".xls"],
            tooltip="Fichier Excel avec les colonnes email et nom des tiers (produit étape 1, complété)",
        ),
        ModuleInput(
            key="dossier_pdf",
            label="Dossier des lettres PDF individuelles",
            input_type="folder",
            tooltip="Dossier contenant les PDF produits par l'étape 3 (lettres/output/)",
        ),
        ModuleInput(
            key="nom_societe",
            label="Nom de la société auditée",
            input_type="text",
        ),
        ModuleInput(
            key="date_arrete",
            label="Date d'arrêté",
            input_type="text",
            default="31/12/2024",
        ),
        ModuleInput(
            key="type_tiers",
            label="Type de tiers",
            input_type="combo",
            options=["Clients", "Fournisseurs"],
            default="Clients",
            required=False,
        ),
    ],
}


class CircularisationTiers(BaseModule):

    name = "Circularisation des Tiers"
    description = (
        "Automatisation NEP 505 / ISA 505 — 4 étapes indépendantes :\n"
        "Sélection → Lettres → Split PDF → Emails"
    )
    category = "Audit"
    version = "3.0"
    help_text = (
        "WORKFLOW EN 4 ÉTAPES INDÉPENDANTES\n\n"
        "1. SÉLECTION : charge la balance auxiliaire et sélectionne les tiers\n"
        "   à circulariser par cumul décroissant (85-95 %).\n\n"
        "2. GÉNÉRATION DES LETTRES : remplit un modèle Word avec les données\n"
        "   de chaque tiers (nom, montant, date de clôture).\n\n"
        "3. SPLIT PDF : découpe le PDF signé en fichiers PDF individuels\n"
        "   nommés par tiers, prêts à être envoyés.\n\n"
        "4. GÉNÉRATION EMAILS : crée des fichiers .eml personnalisés avec\n"
        "   le PDF en pièce jointe — à ouvrir dans Outlook pour envoi.\n\n"
        "Vous pouvez exécuter chaque étape séparément selon votre avancement."
    )
    detection_keywords = ["client", "fournisseur", "tiers", "circularisation",
                          "confirmation", "balance auxiliaire", "solde"]
    detection_threshold = 0.4

    # Étape active (modifiée par l'UI)
    etape_active: str = "selection"

    # ── API BaseModule ─────────────────────────────────────────

    def get_required_inputs(self):
        return INPUTS_PAR_ETAPE.get(self.etape_active, [])

    def get_param_schema(self):
        return []  # Pas de paramètres cachés — tout est dans les inputs

    def validate(self, inputs):
        errors = []
        etape = inputs.get("_etape", self.etape_active)
        required = INPUTS_PAR_ETAPE.get(etape, [])
        for inp in required:
            if not inp.required:
                continue
            val = inputs.get(inp.key, "")
            if not val:
                errors.append(f"'{inp.label}' est requis.")
            elif inp.input_type == "file" and not Path(str(val)).exists():
                errors.append(f"Fichier introuvable : {val}")
            elif inp.input_type == "folder" and not Path(str(val)).exists():
                errors.append(f"Dossier introuvable : {val}")
        return (not errors, errors)

    def preview(self, inputs):
        etape = inputs.get("_etape", self.etape_active)
        # Aperçu du fichier principal selon l'étape
        file_keys = {
            "selection": "balance_aux",
            "lettres":   "fichier_selection",
            "split":     "pdf_signe",
            "emails":    "fichier_selection",
        }
        key = file_keys.get(etape)
        if key:
            p = inputs.get(key, "")
            if p and Path(str(p)).exists() and str(p).endswith((".xlsx", ".xls")):
                try:
                    return pd.read_excel(p, nrows=10)
                except Exception:
                    pass
        return None

    def execute(self, inputs, output_dir, progress_callback=None):
        etape = inputs.get("_etape", self.etape_active)
        os.makedirs(output_dir, exist_ok=True)

        if progress_callback:
            progress_callback(5, f"Démarrage — {etape}...")

        # Charger le moteur script.py
        script_path = Path(__file__).parent / "script.py"
        if not script_path.exists():
            return ModuleResult(
                success=False,
                message="script.py introuvable dans modules/circularisation/\n"
                        "Ajoutez le fichier script.py du projet circularisation."
            )

        spec = importlib.util.spec_from_file_location("circ_engine", str(script_path))
        eng  = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(eng)

        # Surcharger SCRIPT_DIR pour que les sorties aillent dans output_dir
        eng.SCRIPT_DIR = Path(output_dir)
        eng.DIRS = {k: str(Path(output_dir) / v) for k, v in eng.DIRS.items()}

        # Charger la config
        config = eng.charger_config() if hasattr(eng, "charger_config") else eng.DEFAULT_CONFIG
        # Créer les dossiers nécessaires
        if hasattr(eng, "creer_dossiers"):
            # patch SCRIPT_DIR dans le module
            import types
            eng.SCRIPT_DIR = Path(output_dir)
            eng.creer_dossiers()

        captured = io.StringIO()

        try:
            with redirect_stdout(captured):
                if etape == "selection":
                    result_path = self._run_selection(eng, inputs, config, output_dir, progress_callback)
                elif etape == "lettres":
                    result_path = self._run_lettres(eng, inputs, config, output_dir, progress_callback)
                elif etape == "split":
                    result_path = self._run_split(eng, inputs, config, output_dir, progress_callback)
                elif etape == "emails":
                    result_path = self._run_emails(eng, inputs, config, output_dir, progress_callback)
                else:
                    return ModuleResult(success=False, message=f"Étape inconnue : {etape}")

            if progress_callback:
                progress_callback(100, "Terminé !")

            logs = [l.strip() for l in captured.getvalue().splitlines() if l.strip()]
            log_summary = "\n".join(f"• {l}" for l in logs[-6:]) if logs else ""

            rp = Path(str(result_path)) if result_path else None

            if rp and rp.exists():
                # ── Cas : résultat = DOSSIER (split PDF) ──────────
                if rp.is_dir():
                    pdfs = sorted(rp.glob("*.pdf"))
                    nb   = len(pdfs)
                    taille_totale = sum(f.stat().st_size for f in pdfs) // 1024
                    lignes = [f"   {p.name}  ({p.stat().st_size//1024} Ko)" for p in pdfs[:8]]
                    if nb > 8:
                        lignes.append(f"   … et {nb-8} autres")
                    msg = (
                        f"{nb} PDF créé(s) avec succès\n\n"
                        f"Dossier : {rp.name}\n"
                        f"Taille totale : {taille_totale} Ko\n\n"
                        + "\n".join(lignes)
                    )
                    return ModuleResult(
                        success=True,
                        output_path=str(rp),
                        message=msg,
                        stats={"PDFs créés": nb, "Taille (Ko)": taille_totale},
                    )

                # ── Cas : résultat = FICHIER (sélection, lettres, emails) ──
                ext  = rp.suffix.lower()
                size = rp.stat().st_size // 1024
                labels = {
                    ".xlsx": "Fichier Excel",
                    ".docx": "Document Word",
                    ".pdf":  "PDF",
                    ".eml":  "Email",
                }
                type_label = labels.get(ext, "Fichier")
                msg = (
                    f"{type_label} généré avec succès\n\n"
                    f"Fichier : {rp.name}\n"
                    f"Taille  : {size} Ko"
                    + (f"\n\n{log_summary}" if log_summary else "")
                )
                return ModuleResult(
                    success=True,
                    output_path=str(rp),
                    message=msg,
                )

            # ── Cas : aucun résultat explicite → chercher le plus récent ──
            all_items = [f for f in Path(output_dir).rglob("*") if f.is_file()]
            if all_items:
                newest = max(all_items, key=lambda f: f.stat().st_mtime)
                return ModuleResult(
                    success=True,
                    output_path=str(newest),
                    message=f"Traitement terminé.\n\nFichier : {newest.name}"
                            + (f"\n\n{log_summary}" if log_summary else ""),
                )
            return ModuleResult(
                success=False,
                message=f"Aucun fichier produit.\n\n{log_summary}"
            )

        except Exception as e:
            logs = [l.strip() for l in captured.getvalue().splitlines() if l.strip()]
            detail = "\n".join(f"• {l}" for l in logs[-6:])
            return ModuleResult(
                success=False,
                errors=[f"Erreur étape '{etape}' : {e}\n\n{detail}"]
            )

    # ── Runners par étape ──────────────────────────────────────

    def _run_selection(self, eng, inputs, config, output_dir, cb):
        if cb: cb(20, "Chargement du fichier…")
        chemin = inputs["balance_aux"]
        type_tiers = inputs.get("type_tiers", "Fournisseurs").lower()
        type_tiers = "fournisseurs" if "fourn" in type_tiers else "clients"

        seuil_min = float(inputs.get("couverture_min", 85)) / 100
        seuil_max = float(inputs.get("couverture_max", 95)) / 100

        if cb: cb(35, "Analyse du format…")

        if eng.est_format_dynamics(chemin):
            df, societe = eng.parser_dynamics_ax(chemin, type_tiers)
            col_montant = "total_credit" if type_tiers == "fournisseurs" else "solde_closing"
        else:
            df = eng.charger_fichier(chemin)
            if df is None or df.empty:
                raise ValueError("Fichier vide ou illisible.")
            mapping = _auto_mapper(df, type_tiers, eng, config)
            col_tiers   = mapping.get("tiers")
            col_nom     = mapping.get("nom_tiers", col_tiers)
            col_montant = mapping.get("montant_credit" if type_tiers == "fournisseurs" else "solde", None)
            if not col_tiers or not col_montant:
                raise ValueError(
                    f"Colonnes non identifiées automatiquement.\n"
                    f"Colonnes disponibles : {list(df.columns)}\n"
                    f"Renommez-les en : tiers, nom, solde (ou montant_credit)"
                )
            df[col_montant] = pd.to_numeric(df[col_montant], errors="coerce").fillna(0)
            df = df.groupby(col_tiers).agg(
                {col_montant: "sum", **({col_nom: "first"} if col_nom != col_tiers else {})}
            ).reset_index()
            df = df.rename(columns={col_tiers: "code_tiers", col_montant: "montant"})
            if col_nom and col_nom != col_tiers and col_nom in df.columns:
                df = df.rename(columns={col_nom: "nom"})
            else:
                df["nom"] = df["code_tiers"]
            col_montant = "montant"

        if cb: cb(55, "Sélection par cumul décroissant…")

        df_positifs = df[df[col_montant].abs() > 0]
        sel, stats  = eng.selectionner_par_cumul(df_positifs, col_montant, seuil_min, seuil_max)

        if "email" not in sel.columns:   sel["email"]   = ""
        if "adresse" not in sel.columns: sel["adresse"] = ""

        if cb: cb(75, "Export Excel…")

        date_str  = datetime.now().strftime("%Y%m%d_%H%M")
        out_dir   = Path(output_dir)
        out_path  = out_dir / f"selection_{type_tiers}_{date_str}.xlsx"
        sel.to_excel(out_path, index=False, engine="openpyxl")

        _format_excel_selection(out_path)

        return out_path

    def _run_lettres(self, eng, inputs, config, output_dir, cb):
        if cb: cb(15, "Chargement des données…")
        df = pd.read_excel(inputs["fichier_selection"])

        # Déterminer automatiquement le type de tiers depuis le nom du fichier
        filename = Path(inputs["fichier_selection"]).name.lower()
        is_client = "client" in filename or "clt" in filename
        canvas_name = "lc client.docx" if is_client else "lc fournisseur.docx"
        canvas_path = Path(__file__).parent / "canvas" / canvas_name

        if not canvas_path.exists():
            raise FileNotFoundError(f"Canvas introuvable : {canvas_path}")

        if cb: cb(30, f"Utilisation du canvas : {canvas_name}")
        template = eng.analyser_template(str(canvas_path))

        date_lettre = inputs.get("date_lettre", datetime.now().strftime("%d/%m/%Y"))

        # Extraire le nom de société depuis le fichier de sélection (première ligne ou config)
        nom_societe = "SOCIÉTÉ À AUDITER"  # Valeur par défaut
        try:
            # Essayer d'extraire depuis les données
            if "nom_societe" in df.columns and not df.empty:
                nom_societe = str(df["nom_societe"].iloc[0]) if pd.notna(df["nom_societe"].iloc[0]) else nom_societe
        except:
            pass

        config_gen = {**config.get("cabinet", {}),
                      "nom_societe": nom_societe, "date_cloture": date_lettre}

        col_nom   = _find_col(df, ["nom", "name", "nom_tiers", "libellé", "raison sociale"])
        col_code  = _find_col(df, ["code_tiers", "code", "compte", "n°"])
        col_solde = _find_col(df, ["solde", "montant", "closing", "credit", "balance"])

        if cb: cb(40, f"Génération de {len(df)} lettres…")
        out_dir = Path(output_dir)
        fichiers_gen = []

        for _, row in df.iterrows():
            tiers = {
                "nom":          str(row.get(col_nom, row.get(col_code, "Inconnu"))),
                "code_tiers":   str(row.get(col_code, "")),
                "solde_closing": float(row[col_solde]) if col_solde and pd.notna(row.get(col_solde)) else 0,
                "devise":       str(row.get("devise", "MAD")),
            }
            doc      = eng.remplir_template(template["chemin"], template["zones"], tiers, config_gen, date_lettre)
            nom_safe = eng.nettoyer_nom(tiers["nom"])
            nom_fic  = f"{nom_safe}_{tiers['code_tiers']}.docx"
            chemin_out = out_dir / nom_fic
            doc.save(str(chemin_out))
            fichiers_gen.append({"chemin": str(chemin_out), "nom": tiers["nom"], "code": tiers["code_tiers"]})

        if cb: cb(80, "Consolidation…")

        col_nom   = _find_col(df, ["nom", "name", "nom_tiers", "libellé", "raison sociale"])
        col_code  = _find_col(df, ["code_tiers", "code", "compte", "n°"])
        col_solde = _find_col(df, ["solde", "montant", "closing", "credit", "balance"])

        if cb: cb(40, f"Génération de {len(df)} lettres…")
        out_dir = Path(output_dir)
        fichiers_gen = []

        for _, row in df.iterrows():
            tiers = {
                "nom":          str(row.get(col_nom, row.get(col_code, "Inconnu"))),
                "code_tiers":   str(row.get(col_code, "")),
                "solde_closing": float(row[col_solde]) if col_solde and pd.notna(row.get(col_solde)) else 0,
                "devise":       str(row.get("devise", "MAD")),
            }
            doc      = eng.remplir_template(template["chemin"], template["zones"], tiers, config_gen, date_cloture)
            nom_safe = eng.nettoyer_nom(tiers["nom"])
            nom_fic  = f"LC_{tiers['code_tiers']}_{nom_safe}.docx"
            chemin_out = out_dir / nom_fic
            doc.save(str(chemin_out))
            fichiers_gen.append({"chemin": str(chemin_out), "nom": tiers["nom"], "code": tiers["code_tiers"]})

        if cb: cb(80, "Consolidation…")

        # Document consolidé
        if fichiers_gen:
            from docx import Document as DocxDoc
            nom_soc_safe = eng.nettoyer_nom(nom_societe) if nom_societe else "circularisation"
            type_tag = "Frs" if "fourn" in Path(inputs["fichier_selection"]).name.lower() else "Clients"
            nom_cons   = f"Circularisation_{type_tag}_{nom_soc_safe}.docx"
            doc_final  = DocxDoc(fichiers_gen[0]["chemin"])
            for fg in fichiers_gen[1:]:
                doc_final.add_page_break()
                doc_src = DocxDoc(fg["chemin"])
                for para in doc_src.paragraphs:
                    np_ = doc_final.add_paragraph()
                    np_.paragraph_format.alignment = para.paragraph_format.alignment
                    for run in para.runs:
                        nr = np_.add_run(run.text)
                        nr.bold, nr.italic, nr.underline = run.bold, run.italic, run.underline
                        if run.font.size: nr.font.size = run.font.size
                        if run.font.name: nr.font.name = run.font.name
            chemin_cons = out_dir / nom_cons
            doc_final.save(str(chemin_cons))

            import json
            etat = {"mapping_pages": [{"tiers": f["nom"], "code": f["code"]} for f in fichiers_gen],
                    "type": type_tag, "date": datetime.now().isoformat()}
            with open(out_dir / "etat_generation.json", "w", encoding="utf-8") as f:
                json.dump(etat, f, indent=2, ensure_ascii=False)

            return chemin_cons
        return None

    def _run_split(self, eng, inputs, config, output_dir, cb):
        try:
            from pypdf import PdfReader, PdfWriter
        except ImportError:
            raise ImportError(
                "Le module 'pypdf' est requis pour le split PDF.\n"
                "Installez-le : pip install pypdf"
            )
        import json

        if cb: cb(10, "Lecture du PDF…")
        pdf_path       = inputs.get("pdf_signe", "")
        pages_p_lettre = 1  # Champ masqué en UI: valeur standard 1 page/lettre

        if not pdf_path or not Path(str(pdf_path)).exists():
            raise FileNotFoundError(f"PDF introuvable : {pdf_path}")

        reader      = PdfReader(str(pdf_path))
        total_pages = len(reader.pages)

        if total_pages == 0:
            raise ValueError("Le PDF est vide (0 pages).")

        if cb: cb(20, f"PDF chargé : {total_pages} page(s)…")

        # ── Mapping des tiers ────────────────────────────────
        mapping = []

        # 1) Source prioritaire : état de génération des lettres (fiable)
        state_path = _find_generation_state_for_split(output_dir, str(pdf_path))
        if state_path:
            mapping = _load_mapping_from_generation_state(state_path)
            if cb and mapping:
                cb(23, f"Mapping lettres chargé ({len(mapping)} tiers)")

            # Si le mapping dépasse le nombre de lettres, on tronque proprement
            nb_lettres_pdf = max(1, total_pages // pages_p_lettre)
            if len(mapping) > nb_lettres_pdf:
                mapping = mapping[:nb_lettres_pdf]

        sel_path = inputs.get("fichier_selection", "")
        if not sel_path:
            sel_path = _find_selection_file_for_split(output_dir, str(pdf_path))

        if (not mapping) and sel_path and str(sel_path).strip() and Path(str(sel_path)).exists():
            df_sel   = pd.read_excel(sel_path)
            col_nom  = _find_col(df_sel, [
                "nom", "name", "nom_tiers", "libellé", "raison sociale", "intitulé",
                "libelle", "fournisseur", "client", "tiers",
                "denomination", "dénomination", "raison_sociale", "intitule",
            ])
            col_code = _find_col(df_sel, [
                "code_tiers", "code", "compte", "n°", "id", "numero", "num",
                "reference", "ref",
            ])

            # Renforcer la détection du nom tiers par analyse du contenu.
            # On n'écrase PAS un bon match d'en-tête trouvé (sinon risque de prendre nom_societe).
            if col_nom is None:
                col_nom_guess = _guess_best_text_column(df_sel, exclude={col_code} if col_code else None)
                if col_nom_guess is not None:
                    col_nom = col_nom_guess

            # Si la colonne détectée est quasi constante, elle représente souvent la société auditée.
            # Dans ce cas, on tente une autre colonne plus discriminante.
            if col_nom is not None and col_nom in df_sel.columns:
                vals_nom = df_sel[col_nom].dropna().astype(str).str.strip()
                vals_nom = vals_nom[~vals_nom.str.lower().isin(["", "nan", "none"])]
                uniq_ratio = (vals_nom.nunique() / max(len(vals_nom), 1)) if not vals_nom.empty else 0.0
                if uniq_ratio < 0.10:
                    alt_nom = _guess_best_text_column(df_sel, exclude={col_code, col_nom} if col_code else {col_nom})
                    if alt_nom is not None:
                        col_nom = alt_nom

            # Fallback final si aucune colonne exploitable n'a été trouvée
            if col_nom is None and len(df_sel.columns) > 0:
                col_nom = df_sel.columns[0]

            if cb:
                cb(25, f"Colonnes détectées — nom: {col_nom} | code: {col_code or 'aucune'}")

            mapping = _build_mapping_from_columns(df_sel, col_nom, col_code)

            # Si les noms sont quasi constants, on force une alternative pour éviter
            # d'utiliser le nom de la société auditée sur toutes les lettres.
            if _mapping_diversity(mapping) < 0.20:
                alt_nom = _guess_best_text_column(
                    df_sel,
                    exclude={col_code, col_nom} if col_code else {col_nom},
                )
                if alt_nom:
                    alt_mapping = _build_mapping_from_columns(df_sel, alt_nom, col_code)
                    if _mapping_diversity(alt_mapping) > _mapping_diversity(mapping):
                        mapping = alt_mapping

            # Dernier garde-fou: si toujours trop uniforme, ignorer ce mapping
            # et laisser le fallback OCR détecter les tiers lettre par lettre.
            if _mapping_diversity(mapping) < 0.12:
                if cb:
                    cb(26, "Mapping sélection trop uniforme — fallback OCR activé")
                mapping = []

        # Mapping automatique depuis le contenu du PDF signé (nom tiers proche du bloc date)
        if not mapping:
            mapping = _extract_tiers_from_signed_pdf(
                reader,
                pages_p_lettre,
                pdf_path=str(pdf_path),
                cb=cb,
            )
            if cb and mapping:
                cb(28, f"Noms tiers détectés automatiquement depuis PDF : {len(mapping)}")

        # Mapping automatique si aucune sélection fournie
        if not mapping:
            nb_lettres = max(1, total_pages // pages_p_lettre)
            mapping = [{"tiers": f"Lettre_{i+1}", "code": f"{i+1:03d}"}
                       for i in range(nb_lettres)]

        if cb: cb(30, f"Découpage : {total_pages} page(s) → {len(mapping)} lettre(s)…")

        # ── Sous-dossier dédié avec timestamp ─────────────────
        ts      = datetime.now().strftime("%Y%m%d_%H%M%S")
        out_dir = Path(output_dir) / f"split_{ts}"
        out_dir.mkdir(parents=True, exist_ok=True)

        page_courante  = 0
        fichiers_split = []
        noms_utilises  = set()  # éviter les doublons de noms de fichiers

        for i, tiers_info in enumerate(mapping):
            nom  = tiers_info.get("tiers", f"Lettre_{i+1}") or f"Lettre_{i+1}"
            code = tiers_info.get("code", f"{i+1:03d}") or f"{i+1:03d}"

            debut = page_courante
            fin   = min(debut + pages_p_lettre, total_pages)
            if debut >= total_pages:
                break

            writer = PdfWriter()
            for p in range(debut, fin):
                writer.add_page(reader.pages[p])

            nom_safe  = eng.nettoyer_nom(nom)  if nom  else ""
            code_safe = eng.nettoyer_nom(code) if code else ""

            # Nommage : priorité au nom du tiers, fallback sur code, puis index
            if nom_safe and nom_safe not in ("nan", "None"):
                base = nom_safe
            elif code_safe and code_safe not in ("nan", "None"):
                base = code_safe
            else:
                base = f"tiers_{i+1:03d}"

            nom_fic = f"{base}.pdf"
            if nom_fic in noms_utilises:
                nom_fic = f"{base}_{i+1:03d}.pdf"
            noms_utilises.add(nom_fic)

            chemin_out = out_dir / nom_fic
            with open(chemin_out, "wb") as f:
                writer.write(f)

            taille_kb = chemin_out.stat().st_size // 1024
            fichiers_split.append({
                "chemin": str(chemin_out),
                "tiers":  nom,
                "code":   code,
                "pages":  f"{debut+1}-{fin}",
                "taille_kb": taille_kb,
            })
            page_courante = fin

            if cb:
                pct = 30 + int(65 * (i + 1) / len(mapping))
                cb(pct, f"{i+1}/{len(mapping)} — {nom[:35]}")

        # Pages restantes non couvertes
        restantes = total_pages - page_courante
        if restantes > 0:
            writer = PdfWriter()
            for p in range(page_courante, total_pages):
                writer.add_page(reader.pages[p])
            nom_reste = f"Pages_restantes_{page_courante+1}-{total_pages}.pdf"
            with open(out_dir / nom_reste, "wb") as f:
                writer.write(f)
            fichiers_split.append({
                "chemin": str(out_dir / nom_reste),
                "tiers": "Pages restantes",
                "code": "RESTE",
                "pages": f"{page_courante+1}-{total_pages}",
            })

        # Sauvegarder le manifeste JSON dans le sous-dossier
        with open(out_dir / "split_result.json", "w", encoding="utf-8") as f:
            json.dump(fichiers_split, f, indent=2, ensure_ascii=False)

        if cb: cb(98, f"{len(fichiers_split)} PDF créés dans {out_dir.name}")

        # ── Retourner le DOSSIER contenant les PDFs ───────────
        # (pas le JSON — l'UI ouvrira le dossier dans l'explorateur)
        return out_dir

    def _run_emails(self, eng, inputs, config, output_dir, cb):
        import json
        from email.mime.multipart import MIMEMultipart
        from email.mime.text import MIMEText
        from email.mime.application import MIMEApplication
        from email.utils import formatdate

        if cb: cb(10, "Chargement des données…")
        df         = pd.read_excel(inputs["fichier_selection"])
        pdf_dir    = Path(inputs["dossier_pdf"])
        nom_societe = inputs.get("nom_societe", "")
        date_arrete = inputs.get("date_arrete", "31/12/2024")
        type_tiers  = inputs.get("type_tiers", "Clients")
        nom_cabinet = config.get("cabinet", {}).get("nom", "AuditPro Cabinet")
        email_exp   = config.get("cabinet", {}).get("email_expediteur", "audit@auditpro.local")

        col_nom   = _find_col(df, ["nom", "name", "nom_tiers", "libellé"])
        col_code  = _find_col(df, ["code_tiers", "code", "compte"])
        col_email = _find_col(df, ["email", "mail", "e-mail", "courriel"])
        col_solde = _find_col(df, ["solde", "montant", "closing", "balance"])

        SUJET = f"Circularisation – Confirmation de solde – {nom_societe}"
        CORPS = (
            f"Bonjour,\n\n"
            f"Dans le cadre de notre mission de commissariat aux comptes pour la société "
            f"« {nom_societe} » au titre de l'exercice clos le {date_arrete}, nous vous "
            f"prions de bien vouloir répondre à la lettre de circularisation ci-jointe.\n\n"
            f"En vous remerciant par avance pour votre retour,\n\n"
            f"Bien cordialement,\n{nom_cabinet}"
        )

        out_dir = Path(output_dir)
        generes, manquants = 0, []

        if cb: cb(20, f"Génération de {len(df)} emails…")

        for i, row in df.iterrows():
            nom   = str(row.get(col_nom, row.get(col_code, f"Tiers_{i}")))
            code  = str(row.get(col_code, ""))
            email = str(row.get(col_email, "")) if col_email else ""
            montant = float(row[col_solde]) if col_solde and pd.notna(row.get(col_solde)) else 0

            nom_safe = eng.nettoyer_nom(nom)
            pdf_candidates = (
                list(pdf_dir.glob(f"*{nom_safe}*.pdf")) +
                list(pdf_dir.glob(f"*{code}*.pdf"))
            )
            pdf_path = pdf_candidates[0] if pdf_candidates else None

            if not pdf_path:
                manquants.append({"tiers": nom, "motif": "PDF manquant"})
                continue
            if not email or "@" not in email:
                manquants.append({"tiers": nom, "motif": "Email manquant"})
                email = ""

            msg = MIMEMultipart()
            msg["Subject"] = SUJET
            msg["To"]      = email
            msg["From"]    = email_exp
            msg["Date"]    = formatdate(localtime=True)
            msg.attach(MIMEText(CORPS, "plain", "utf-8"))

            with open(pdf_path, "rb") as f:
                pj = MIMEApplication(f.read(), _subtype="pdf")
                pj.add_header("Content-Disposition", "attachment", filename=pdf_path.name)
                msg.attach(pj)

            eml_name = f"Email_{code}_{nom_safe}.eml"
            with open(out_dir / eml_name, "w", encoding="utf-8") as f:
                f.write(msg.as_string())
            generes += 1

            if cb: cb(20 + int(70 * (i + 1) / len(df)), f"{i+1}/{len(df)} — {nom[:30]}")

        if manquants:
            rapport = out_dir / "emails_manquants.xlsx"
            pd.DataFrame(manquants).to_excel(rapport, index=False)

        rapport_final = out_dir / f"rapport_emails_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
        pd.DataFrame({
            "Statut":  ["Générés", "Avec problème"],
            "Nombre":  [generes, len(manquants)],
        }).to_excel(rapport_final, index=False)

        return rapport_final


# ── Helpers internes ───────────────────────────────────────────

def _find_col(df, keywords):
    """Trouve la première colonne correspondant à un mot-clé (insensible casse)."""
    for kw in keywords:
        for col in df.columns:
            if kw.lower() in str(col).lower():
                return col
    return None


def _guess_best_text_column(df, exclude=None):
    """Devine la meilleure colonne texte pour représenter le nom du tiers."""
    exclude = set(exclude or [])
    best_col = None
    best_score = -1.0

    forbidden_headers = [
        "nom_societe", "societe_auditee", "société_auditée",
        "societe audit", "société audit", "entite_auditee", "entité_auditée",
    ]

    for col in df.columns:
        if col in exclude:
            continue

        s = df[col]
        # On travaille sur les valeurs non nulles converties en string
        vals = s.dropna().astype(str).str.strip()
        if vals.empty:
            continue

        vals = vals[~vals.str.lower().isin(["", "nan", "none"])]
        if vals.empty:
            continue

        header = str(col).lower()
        if any(k in header for k in forbidden_headers):
            continue

        header_bonus = 0.0
        if any(k in header for k in ["nom", "name", "tiers", "raison", "denomination", "intitule", "intitul"]):
            header_bonus += 2.0
        if any(k in header for k in ["email", "mail", "adresse", "phone", "tel", "ice", "if", "rc"]):
            header_bonus -= 2.0

        # Ratio de valeurs contenant au moins une lettre
        alpha_ratio = vals.str.contains(r"[A-Za-zÀ-ÿ]", regex=True).mean()
        # Une colonne nom est en général variée et pas purement numérique
        uniq_ratio = min(1.0, vals.nunique(dropna=True) / max(len(vals), 1))
        digit_ratio = vals.str.fullmatch(r"\d+").mean()
        dominant_ratio = vals.value_counts(normalize=True).iloc[0] if not vals.empty else 1.0

        score = header_bonus + (2.0 * alpha_ratio) + (0.8 * uniq_ratio) - (1.2 * digit_ratio) - (1.5 * dominant_ratio)
        if score > best_score:
            best_score = score
            best_col = col

    return best_col


def _clean_line_for_name(line: str) -> str:
    line = str(line or "").strip()
    line = re.sub(r"\s+", " ", line)
    return line.strip("-_:;,. ")


def _is_date_like(line: str) -> bool:
    t = line.lower().strip()
    if re.search(r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b", t):
        return True
    if re.search(r"\b(le\s+)?\d{1,2}\s+[a-zà-ÿ]{3,}\s+\d{4}\b", t):
        return True
    return False


def _score_name_line(line: str) -> float:
    t = _clean_line_for_name(line)
    if not t:
        return -999.0

    low = t.lower()
    # Rejeter les lignes de corps ou libellés non nom-tiers
    banned = [
        "messieurs", "madame", "objet", "cachet", "signature", "cabinet",
        "commissaire", "compte", "relevé", "montant", "échéance", "fax", "tel",
        "boulevard", "casablanca", "le ",
    ]
    if any(b in low for b in banned):
        return -999.0
    if _is_date_like(t):
        return -999.0

    letters = re.findall(r"[A-Za-zÀ-ÿ]", t)
    if not letters:
        return -999.0

    alpha_ratio = len(letters) / max(1, len(t))
    upper_ratio = sum(1 for c in letters if c.isupper()) / max(1, len(letters))
    words = [w for w in re.split(r"\s+", t) if w]
    word_bonus = 0.6 if 1 <= len(words) <= 6 else 0.1
    len_bonus = 0.5 if 4 <= len(t) <= 60 else 0.0

    return (2.0 * alpha_ratio) + (1.2 * upper_ratio) + word_bonus + len_bonus


def _extract_tiers_from_signed_pdf(
    reader,
    pages_p_lettre: int,
    pdf_path: str | None = None,
    cb=None,
):
    """Extrait des noms de tiers depuis le PDF signé, autour du bloc date en haut de page."""
    mapping = []
    total_pages = len(reader.pages)
    if total_pages <= 0:
        return mapping

    # Garde-fou: éviter un blocage long sur gros PDF scannés.
    max_ocr_pages = 40
    ocr_used = 0

    i = 0
    while i < total_pages:
        if cb:
            pct = 21 + int(6 * (i / max(total_pages, 1)))
            cb(pct, f"Analyse des lettres ({min(i + 1, total_pages)}/{total_pages})…")

        page = reader.pages[i]
        text = page.extract_text() or ""
        if not text.strip() and pdf_path and ocr_used < max_ocr_pages:
            text = _extract_page_text_with_ocr(pdf_path, i)
            ocr_used += 1

        lines = [_clean_line_for_name(l) for l in text.splitlines()]
        lines = [l for l in lines if l]

        chosen = None
        if lines:
            # On privilégie les candidats autour d'une ligne date détectée
            date_idxs = [idx for idx, ln in enumerate(lines[:20]) if _is_date_like(ln)]
            candidate_pool = []
            if date_idxs:
                d0 = date_idxs[0]
                lo = max(0, d0 - 3)
                hi = min(len(lines), d0 + 4)
                candidate_pool = lines[lo:hi]
            else:
                # Fallback : top de page uniquement
                candidate_pool = lines[:12]

            best = None
            best_score = -999.0
            for ln in candidate_pool:
                sc = _score_name_line(ln)
                if sc > best_score:
                    best_score = sc
                    best = ln

            if best and best_score > 0:
                chosen = best

        if not chosen:
            chosen = f"Lettre_{len(mapping)+1}"

        mapping.append({"tiers": chosen, "code": f"{len(mapping)+1:03d}"})
        i += max(1, pages_p_lettre)

    return mapping


def _load_ocr_paths_for_split():
    """Charge core/ocr_paths.py pour récupérer les chemins Tesseract/Poppler."""
    mod_name = "core.ocr_paths"
    if mod_name in sys.modules:
        return sys.modules[mod_name]

    here = Path(__file__).resolve()
    for parent in here.parents:
        p = parent / "core" / "ocr_paths.py"
        if p.exists():
            spec = importlib.util.spec_from_file_location(mod_name, str(p))
            if spec and spec.loader:
                mod = importlib.util.module_from_spec(spec)
                sys.modules[mod_name] = mod
                spec.loader.exec_module(mod)
                return mod
    return None


def _extract_page_text_with_ocr(pdf_path: str, page_index: int) -> str:
    """OCR d'une page précise d'un PDF signé quand extract_text() est vide."""
    try:
        import pytesseract
        from pdf2image import convert_from_path
    except Exception:
        return ""

    poppler_path = None
    try:
        ocr_mod = _load_ocr_paths_for_split()
        if ocr_mod and getattr(ocr_mod, "TESSERACT_PATH", None):
            pytesseract.pytesseract.tesseract_cmd = str(ocr_mod.TESSERACT_PATH)
        if ocr_mod and getattr(ocr_mod, "POPPLER_PATH", None):
            poppler_path = str(ocr_mod.POPPLER_PATH)
    except Exception:
        pass

    try:
        kwargs = {
            "first_page": page_index + 1,
            "last_page": page_index + 1,
            "dpi": 170,
        }
        if poppler_path:
            kwargs["poppler_path"] = poppler_path

        images = convert_from_path(pdf_path, **kwargs)
        if not images:
            return ""

        # fra+eng pour couvrir les modèles FR avec noms internationaux
        txt = pytesseract.image_to_string(images[0], lang="fra+eng", timeout=8) or ""
        try:
            images[0].close()
        except Exception:
            pass
        return txt
    except Exception:
        return ""


def _auto_mapper(df, type_tiers, eng, config):
    """Mapping automatique sans interaction (utilise les keywords du script)."""
    rules = config.get("mapping", {}).get(type_tiers, {})
    mapping = {}
    for role, keywords in rules.items():
        col = eng.trouver_colonne(df, keywords)
        if col:
            mapping[role] = col
    return mapping


def _find_selection_file_for_split(output_dir: str, pdf_path: str) -> str:
    """Trouve automatiquement le dernier fichier de sélection Excel pour l'étape split."""
    candidates = []
    roots = []

    try:
        outp = Path(output_dir)
        roots.extend([outp, outp.parent])
    except Exception:
        pass

    try:
        pdfp = Path(pdf_path)
        roots.extend([pdfp.parent, pdfp.parent.parent])
    except Exception:
        pass

    seen = set()
    uniq_roots = []
    for r in roots:
        if not r:
            continue
        try:
            rr = r.resolve()
        except Exception:
            rr = r
        if rr in seen or not rr.exists() or not rr.is_dir():
            continue
        seen.add(rr)
        uniq_roots.append(rr)

    patterns = ["selection_*.xlsx", "selection_*.xls", "*selection*.xlsx", "*selection*.xls"]
    for root in uniq_roots:
        for pat in patterns:
            for p in root.rglob(pat):
                if p.is_file() and not p.name.startswith("~$"):
                    candidates.append(p)

    if not candidates:
        return ""

    latest = max(candidates, key=lambda p: p.stat().st_mtime)
    return str(latest)


def _find_generation_state_for_split(output_dir: str, pdf_path: str) -> str:
    """Trouve le dernier etat_generation.json issu de l'étape lettres."""
    candidates = []
    roots = []

    try:
        outp = Path(output_dir)
        roots.extend([outp, outp.parent])
    except Exception:
        pass

    try:
        pdfp = Path(pdf_path)
        roots.extend([pdfp.parent, pdfp.parent.parent])
    except Exception:
        pass

    seen = set()
    uniq_roots = []
    for r in roots:
        if not r:
            continue
        try:
            rr = r.resolve()
        except Exception:
            rr = r
        if rr in seen or not rr.exists() or not rr.is_dir():
            continue
        seen.add(rr)
        uniq_roots.append(rr)

    for root in uniq_roots:
        for p in root.rglob("etat_generation.json"):
            if p.is_file():
                candidates.append(p)

    if not candidates:
        return ""
    latest = max(candidates, key=lambda p: p.stat().st_mtime)
    return str(latest)


def _load_mapping_from_generation_state(state_path: str):
    """Charge mapping_pages depuis etat_generation.json en format split mapping."""
    try:
        import json
        with open(state_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        rows = data.get("mapping_pages", []) if isinstance(data, dict) else []
        mapping = []
        for r in rows:
            tiers = str((r or {}).get("tiers", "")).strip()
            code = str((r or {}).get("code", "")).strip()
            if tiers and tiers.lower() not in {"nan", "none"}:
                mapping.append({"tiers": tiers, "code": code})
        return mapping
    except Exception:
        return []


def _build_mapping_from_columns(df_sel, col_nom, col_code):
    mapping = []
    for _, row in df_sel.iterrows():
        nom_val = str(row[col_nom]).strip() if col_nom and pd.notna(row.get(col_nom)) else ""
        code_val = str(row[col_code]).strip() if col_code and pd.notna(row.get(col_code)) else ""
        nom_val = "" if nom_val in ("nan", "None") else nom_val
        code_val = "" if code_val in ("nan", "None") else code_val
        if nom_val or code_val:
            mapping.append({"tiers": nom_val or code_val, "code": code_val})
    return mapping


def _mapping_diversity(mapping):
    """Retourne un ratio [0..1] de diversité des noms tiers dans le mapping."""
    if not mapping:
        return 0.0
    names = [str(x.get("tiers", "")).strip().lower() for x in mapping]
    names = [n for n in names if n and n not in ("nan", "none")]
    if not names:
        return 0.0
    return len(set(names)) / max(len(names), 1)


def _format_excel_selection(path: Path):
    """Mise en forme Excel de la sélection."""
    try:
        from openpyxl import load_workbook
        from openpyxl.styles import Font, PatternFill, Border, Side
        wb = load_workbook(path)
        ws = wb.active
        hdr_fill = PatternFill("solid", fgColor="4B286D")   # violet GT
        hdr_font = Font(bold=True, color="FFFFFF", name="Arial", size=10)
        vert_fill = PatternFill("solid", fgColor="EDE7F6")
        bd = Border(left=Side("thin"), right=Side("thin"),
                    top=Side("thin"), bottom=Side("thin"))
        for cell in ws[1]:
            cell.fill, cell.font, cell.border = hdr_fill, hdr_font, bd
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row, max_col=ws.max_column):
            for cell in row:
                cell.border = bd
                cell.font   = Font(name="Arial", size=10)
                hdr = str(ws.cell(1, cell.column).value or "").lower()
                if hdr in ("email", "adresse") and not cell.value:
                    cell.fill = vert_fill
        ws.freeze_panes = "A2"
        for col in ws.columns:
            ws.column_dimensions[col[0].column_letter].width = max(
                12, min(42, max(len(str(c.value or "")) for c in col) + 2)
            )
        wb.save(path)
    except Exception:
        pass  # Mise en forme optionnelle
