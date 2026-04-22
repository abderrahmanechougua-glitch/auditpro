#!/usr/bin/env python3
"""
╔══════════════════════════════════════════════════════════════════╗
║  AUTOMATISATION CIRCULARISATION - AUDIT                        ║
║  Conforme NEP 505 — 100% Python — Windows                      ║
║  v2.0 — Script unique autonome                                 ║
╚══════════════════════════════════════════════════════════════════╝
"""
import os, sys, json, re, subprocess, logging, shutil
from pathlib import Path
from datetime import datetime
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from email.utils import formatdate

# === AUTO-INSTALLATION ===
SCRIPT_DIR = Path(__file__).resolve().parent
os.chdir(SCRIPT_DIR)

_DEPS = {"pandas":"pandas","openpyxl":"openpyxl","docx":"python-docx","pypdf":"pypdf","pdfplumber":"pdfplumber"}
_miss = []
for _i,_p in _DEPS.items():
    try: __import__(_i)
    except ImportError: _miss.append(_p)
if _miss:
    print(f"  Installation: {_miss}")
    subprocess.check_call([sys.executable,"-m","pip","install","--quiet"]+_miss)
    print("  ✓ OK\n")

import pandas as pd
from docx import Document as DocxDocument
from pypdf import PdfReader, PdfWriter

# ═══════════════════════════════════════════════════════════════
#  CONFIGURATION & CONSTANTES
# ═══════════════════════════════════════════════════════════════

DEFAULT_CONFIG = {
    "cabinet": {
        "nom": "AuditPro Cabinet",
        "email_expediteur": "audit@auditpro.local",
        "date_arretes_defaut": "31/12/2025"
    },
    "mapping": {
        "fournisseurs": {
            "tiers": ["Code fournisseur","N°","Code","Compte","Fournisseur","Tiers","Nom",
                       "Supplier account","Vendor account","Name"],
            "nom_tiers": ["Intitulé","Libellé","Nom","Name","Raison sociale","Description",
                          "Désignation","Company","Account Name"],
            "montant_credit": ["Mouvement créditeur","Crédit","Credit","Mvt crédit",
                               "Montant crédit","Créditeur","Credit amount","Amount"],
            "solde": ["Solde","Balance","Solde final","Closing balance","Net balance"],
            "email": ["Email","Mail","Adresse mail","E-mail","Courriel"]
        },
        "clients": {
            "tiers": ["Code client","N°","Code","Compte","Client","Tiers","Nom",
                       "Customer account","Name"],
            "nom_tiers": ["Intitulé","Libellé","Nom","Name","Raison sociale","Description",
                          "Account Name"],
            "solde": ["Solde","Balance","Montant","Solde final","Closing balance","Amount"],
            "email": ["Email","Mail","Adresse mail","E-mail","Courriel"]
        }
    },
    "seuils_selection": {
        "fournisseurs_intervalle_defaut": "85-95",
        "clients_pourcentage_defaut": "85-95"
    }
}

DIRS = {
    "sel_input_clients":    "selection/input/clients",
    "sel_input_frs":        "selection/input/fournisseurs",
    "sel_output":           "selection/output",
    "sel_canvas":           "selection/canvas",
    "lettres_input":        "lettres/input",
    "lettres_output":       "lettres/output",
    "lettres_temp":         "lettres/temp",
    "emails_output":        "emails/output",
    "emails_templates":     "emails/templates",
    "emails_logs":          "emails/logs",
}

MOIS_FR = {1:'janvier',2:'février',3:'mars',4:'avril',5:'mai',6:'juin',
           7:'juillet',8:'août',9:'septembre',10:'octobre',11:'novembre',12:'décembre'}

# ═══════════════════════════════════════════════════════════════
#  LOGGING
# ═══════════════════════════════════════════════════════════════

def setup_logging():
    log_dir = SCRIPT_DIR / "emails" / "logs"
    log_dir.mkdir(parents=True, exist_ok=True)
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(log_dir / "execution.log", encoding='utf-8'),
            logging.StreamHandler(sys.stdout)
        ]
    )
    # Reduce pandas/other noise
    logging.getLogger('pandas').setLevel(logging.WARNING)
    logging.getLogger('openpyxl').setLevel(logging.WARNING)

# ═══════════════════════════════════════════════════════════════
#  UTILITAIRES
# ═══════════════════════════════════════════════════════════════

def creer_dossiers():
    for nom, chemin in DIRS.items():
        (SCRIPT_DIR / chemin).mkdir(parents=True, exist_ok=True)

def charger_config():
    chemin = SCRIPT_DIR / "config.json"
    if not chemin.exists():
        with open(chemin, 'w', encoding='utf-8') as f:
            json.dump(DEFAULT_CONFIG, f, indent=2, ensure_ascii=False)
        print("  ✓ config.json créé avec valeurs par défaut")
    with open(chemin, 'r', encoding='utf-8') as f:
        return json.load(f)

def nettoyer_nom(nom):
    nom = re.sub(r'[<>:"/\\|?*]', '', str(nom))
    return nom.replace(' ','_').strip('_.') [:60]

def formater_montant(montant, devise="MAD"):
    if not montant or montant == 0: return f"{devise} 0,00"
    signe = "-" if montant < 0 else ""
    abs_m = abs(montant)
    partie_ent = int(abs_m)
    partie_dec = round((abs_m - partie_ent)*100)
    s = f"{partie_ent:,}".replace(",", " ")
    return f"{signe}{devise} {s},{partie_dec:02d}"

def formater_date_fr(date_str=None, style="long"):
    if date_str:
        for fmt in ('%d/%m/%Y','%Y-%m-%d'):
            try:
                dt = datetime.strptime(date_str, fmt); break
            except ValueError: continue
        else: dt = datetime.now()
    else: dt = datetime.now()
    if style=="long": return f"{dt.day} {MOIS_FR[dt.month]} {dt.year}"
    elif style=="lettre": return f"Le {dt.day:02d} {MOIS_FR[dt.month]} {dt.year}"
    return f"{dt.day:02d}/{dt.month:02d}/{dt.year}"

def lister_fichiers(dossier, extensions=None):
    if not extensions: extensions = ['.xlsx','.xlsb','.csv','.pdf']
    dossier = Path(dossier)
    if not dossier.exists(): return []
    return [f for f in sorted(dossier.iterdir())
            if f.is_file() and f.suffix.lower() in extensions
            and not f.name.startswith('~') and not f.name.startswith('$')]

def choisir_fichier(dossier, extensions=None, message="Fichiers disponibles"):
    fichiers = lister_fichiers(dossier, extensions)
    if not fichiers:
        print(f"  ⚠ Aucun fichier trouvé dans {dossier}")
        return None
    print(f"\n  {message}:")
    for i, f in enumerate(fichiers, 1):
        taille = f.stat().st_size / 1024
        print(f"    [{i}] {f.name} ({taille:.0f} KB)")
    try:
        choix = input(f"\n  Votre choix (1-{len(fichiers)}): ").strip()
        return fichiers[int(choix)-1]
    except (ValueError, IndexError):
        print("  Choix invalide")
        return None


# ═══════════════════════════════════════════════════════════════
#  CHARGEMENT DES DONNÉES (multi-format)
# ═══════════════════════════════════════════════════════════════

def charger_fichier(chemin):
    """Charge un fichier Excel/CSV/XLSB en DataFrame. Gère multi-onglets."""
    chemin = Path(chemin)
    ext = chemin.suffix.lower()
    try:
        if ext == '.csv':
            for enc in ['utf-8','latin1','cp1252','iso-8859-1']:
                for sep in [';',',','\t','|']:
                    try:
                        df = pd.read_csv(chemin, encoding=enc, sep=sep, nrows=5)
                        if df.shape[1] > 1:
                            print(f"  ✓ CSV chargé (encodage={enc}, sep='{sep}')")
                            return pd.read_csv(chemin, encoding=enc, sep=sep)
                    except: pass
            print("  ✘ Impossible de charger le CSV")
            return None
        elif ext in ('.xlsx', '.xlsb', '.xls'):
            xl = pd.ExcelFile(chemin)
            onglets = xl.sheet_names
            if len(onglets) > 1:
                print(f"\n  {len(onglets)} onglets détectés:")
                for i, o in enumerate(onglets, 1):
                    try:
                        apercu = pd.read_excel(chemin, sheet_name=o, nrows=3)
                        print(f"    [{i}] {o} ({apercu.shape[1]} colonnes)")
                    except: print(f"    [{i}] {o} (illisible)")
                try:
                    choix = input(f"  Onglet (1-{len(onglets)}) [1]: ").strip() or "1"
                    onglet = onglets[int(choix)-1]
                except: onglet = onglets[0]
            else:
                onglet = onglets[0]
            return pd.read_excel(chemin, sheet_name=onglet)
        elif ext == '.pdf':
            return charger_pdf_tableau(chemin)
    except Exception as e:
        print(f"  ✘ Erreur chargement: {e}")
    return None

def charger_pdf_tableau(chemin):
    """Essaie d'extraire un tableau depuis un PDF."""
    try:
        import pdfplumber
        with pdfplumber.open(chemin) as pdf:
            for page in pdf.pages:
                table = page.extract_table()
                if table and len(table) > 1:
                    print(f"  ✓ Tableau extrait du PDF ({len(table)} lignes)")
                    return pd.DataFrame(table[1:], columns=table[0])
    except: pass
    print("  ✘ Impossible d'extraire un tableau du PDF")
    print("    Convertissez le PDF en Excel et réessayez")
    return None

# ═══════════════════════════════════════════════════════════════
#  MAPPING INTELLIGENT DES COLONNES
# ═══════════════════════════════════════════════════════════════

def similarite(a, b):
    """Calcul de similarité simple entre deux chaînes (sans dépendances)."""
    a, b = a.lower().strip(), b.lower().strip()
    if a == b: return 1.0
    if a in b or b in a: return 0.85
    # Jaccard sur les mots
    mots_a, mots_b = set(a.split()), set(b.split())
    if not mots_a or not mots_b: return 0.0
    inter = len(mots_a & mots_b)
    union = len(mots_a | mots_b)
    return inter / union if union else 0.0

def trouver_colonne(df, mots_cles, exclusions=None):
    """Trouve la meilleure colonne correspondant à une liste de mots-clés."""
    if exclusions is None: exclusions = []
    colonnes = list(df.columns)
    meilleur_score, meilleure_col = 0, None

    for col in colonnes:
        col_str = str(col).lower().strip()
        # Exclure
        if any(excl.lower() in col_str for excl in exclusions): continue
        for mot in mots_cles:
            score = similarite(mot, str(col))
            if score > meilleur_score:
                meilleur_score = score
                meilleure_col = col
    return meilleure_col if meilleur_score >= 0.5 else None

def mapper_colonnes(df, type_tiers, config):
    """
    Identifie automatiquement les colonnes du fichier.
    Si échec → demande interactivement.
    Retourne un dict: {role: nom_colonne}
    """
    mapping_rules = config.get('mapping', {}).get(type_tiers, {})
    mapping = {}
    colonnes_display = list(df.columns)

    roles = ['tiers', 'nom_tiers', 'solde', 'email']
    if type_tiers == 'fournisseurs':
        roles = ['tiers', 'nom_tiers', 'montant_credit', 'solde', 'email']

    for role in roles:
        mots_cles = mapping_rules.get(role, [])
        col = trouver_colonne(df, mots_cles) if mots_cles else None
        if col: mapping[role] = col

    # Afficher le résultat
    print(f"\n  === MAPPING DES COLONNES ({type_tiers}) ===")
    print(f"  Colonnes du fichier: {colonnes_display}")
    for role in roles:
        col = mapping.get(role, '❌ NON TROUVÉ')
        status = "✓" if role in mapping else "✘"
        print(f"    {status} {role:<16} → {col}")

    # Vérifier les colonnes essentielles
    essentiels = ['tiers'] if type_tiers == 'clients' else ['tiers']
    manquants = [r for r in essentiels if r not in mapping]

    if manquants:
        print(f"\n  ⚠ Colonnes manquantes: {manquants}")
        print(f"  Colonnes disponibles:")
        for i, c in enumerate(colonnes_display, 1):
            print(f"    [{i}] {c}")
        for role in manquants:
            try:
                choix = input(f"  Colonne pour '{role}' (numéro): ").strip()
                mapping[role] = colonnes_display[int(choix)-1]
            except: pass

    # Demander confirmation
    ok = input("\n  Mapping correct ? (O/n): ").strip().lower()
    if ok == 'n':
        print("  Colonnes disponibles:")
        for i, c in enumerate(colonnes_display, 1):
            print(f"    [{i}] {c}")
        for role in roles:
            try:
                choix = input(f"  Colonne pour '{role}' (numéro, Entrée=ignorer): ").strip()
                if choix: mapping[role] = colonnes_display[int(choix)-1]
            except: pass

    return mapping

# ═══════════════════════════════════════════════════════════════
#  PARSER DYNAMICS AX (format éprouvé)
# ═══════════════════════════════════════════════════════════════

def parser_dynamics_ax(chemin, type_tiers="clients"):
    """Parse Internal/Supplier Account Statement de Dynamics AX."""
    df = pd.read_excel(chemin, header=None)
    nb_cols = len(df.columns)

    def norm_label(val):
        return str(val).strip().rstrip(':').strip() if pd.notna(val) else ""

    def find_val(row_idx, label):
        for c in range(nb_cols):
            v = df.iloc[row_idx, c] if pd.notna(df.iloc[row_idx, c]) else None
            if v and norm_label(v) == label:
                for c2 in range(c+1, nb_cols):
                    v2 = df.iloc[row_idx, c2] if pd.notna(df.iloc[row_idx, c2]) else None
                    if v2 and str(v2).strip() not in ('','nan','NaN'):
                        return str(v2).strip()
        return None

    # Trouver blocs "Name"
    name_idx = []
    for i in range(len(df)):
        for c in range(min(5, nb_cols)):
            v = df.iloc[i, c] if pd.notna(df.iloc[i, c]) else None
            if v and norm_label(v) == "Name":
                name_idx.append(i); break

    if not name_idx: return None, None

    # Détecter société
    societe = None
    for i in range(min(8, len(df))):
        v = str(df.iloc[i, 0] if pd.notna(df.iloc[i, 0]) else "").strip()
        if v and v not in ('nan','Customer - internal account statement',
                           'Vendor - internal account statement','Supplier account statement'):
            if len(v) > 3: societe = v; break
        v1 = str(df.iloc[i, 1] if pd.notna(df.iloc[i, 1]) else "").strip()
        if v1 and v1 not in ('nan',) and len(v1) > 5 and 'statement' not in v1.lower() and 'date' not in v1.lower():
            societe = v1; break

    tiers_list = []
    for bi, start in enumerate(name_idx):
        end = name_idx[bi+1] if bi+1 < len(name_idx) else len(df)
        t = {'nom':'','adresse':'','code_tiers':'','devise':'MAD',
             'solde_closing':0,'total_debit':0,'total_credit':0,'nb_transactions':0}

        for i in range(start, min(end, start+25)):
            if not t['nom']:
                v = find_val(i, "Name")
                if v: t['nom'] = v
            if not t['adresse']:
                v = find_val(i, "Address")
                if v: t['adresse'] = v
            if not t['code_tiers']:
                for lab in ("Customer account","Vendor account","Supplier account"):
                    v = find_val(i, lab)
                    if v: t['code_tiers'] = v; break
            v = find_val(i, "Currency")
            if v: t['devise'] = v

        # Trouver colonne Amount
        amount_col = None
        for i in range(start, min(end, start+30)):
            for c in range(nb_cols):
                v = df.iloc[i, c] if pd.notna(df.iloc[i, c]) else None
                if v and str(v).strip() == "Amount":
                    amount_col = c; break
            if amount_col: break

        for i in range(start, end):
            vals = {norm_label(df.iloc[i,c]):c for c in range(nb_cols) if pd.notna(df.iloc[i,c])}
            if "Closing" in vals:
                cc = amount_col or vals["Closing"]+1
                for c in range(cc, min(cc+5, nb_cols)):
                    v = df.iloc[i,c]
                    if pd.notna(v):
                        try: t['solde_closing'] = float(v); break
                        except: pass
            if any(k in vals for k in ("Opening","Closing","Description","Amount")): continue
            if amount_col:
                v = df.iloc[i, amount_col] if pd.notna(df.iloc[i, amount_col]) else None
                if v:
                    try:
                        m = float(v)
                        if abs(m) > 0.001:
                            if m > 0: t['total_debit'] += m
                            else: t['total_credit'] += abs(m)
                            t['nb_transactions'] += 1
                    except: pass
        tiers_list.append(t)

    result = pd.DataFrame(tiers_list)
    if 'nom' in result.columns and len(result) > 0:
        result = result[result['nom'].astype(str).str.strip() != '']
    if 'code_tiers' in result.columns and len(result) > 0:
        result = result[result['code_tiers'].astype(str).str.strip() != '']
    return result, societe

def est_format_dynamics(chemin):
    """Détecte si un fichier est au format Dynamics AX."""
    try:
        df = pd.read_excel(chemin, header=None, nrows=15)
        for i in range(min(15, len(df))):
            for c in range(min(5, len(df.columns))):
                v = str(df.iloc[i,c]).strip().rstrip(':').strip() if pd.notna(df.iloc[i,c]) else ""
                if v in ("Name","Customer account","Vendor account","Supplier account"):
                    return True
        texte = ' '.join(str(df.iloc[i,c]) for i in range(min(5,len(df))) for c in range(min(3,len(df.columns))) if pd.notna(df.iloc[i,c]))
        if any(kw in texte for kw in ['internal account statement','Supplier account statement']):
            return True
    except: pass
    return False


# ═══════════════════════════════════════════════════════════════
#  PHASE 1 : SÉLECTION DES TIERS
# ═══════════════════════════════════════════════════════════════

def selectionner_par_cumul(df, col_montant, seuil_min, seuil_max):
    """Sélection par cumul décroissant. Retourne (sélectionnés, stats)."""
    if df.empty: return df, {}
    df = df.copy()
    df['_abs'] = df[col_montant].abs()
    df = df.sort_values('_abs', ascending=False).reset_index(drop=True)
    total = df['_abs'].sum()
    if total == 0: return df.head(0), {'total':0,'couverture':0}
    df['_cumul'] = df['_abs'].cumsum()
    df['_couv'] = df['_cumul'] / total

    montant_min = total * seuil_min
    montant_max = total * seuil_max

    # Trouver le point de coupure
    idx = 0
    cumul = 0
    for i, row in df.iterrows():
        cumul += row['_abs']
        idx = i
        if cumul >= montant_min:
            # Vérifier si on dépasse max
            if cumul <= montant_max:
                break
            else:
                # On a dépassé: si c'est le premier, on le garde quand même
                if i == 0: break
                # Sinon on revient au précédent
                idx = i - 1
                cumul -= row['_abs']
                break

    nb = idx + 1
    couv = df.loc[idx, '_couv']
    sel = df.iloc[:nb].drop(columns=['_abs','_cumul','_couv']).copy()
    sel['motif_selection'] = 'CUMUL'
    sel['pct_du_total'] = (sel[col_montant].abs() / total * 100).round(1)
    sel['cumul_pct'] = (sel[col_montant].abs().cumsum() / total * 100).round(1)

    stats = {'total_pop':len(df), 'nb_sel':nb, 'total_montant':total,
             'couverture':couv, 'montant_couvert':cumul}
    return sel, stats

def phase1_selection(config):
    """Phase interactive de sélection des tiers."""
    print("\n" + "="*60)
    print("  PHASE 1 : SÉLECTION DES TIERS")
    print("="*60)

    # Choix client ou fournisseur
    print("\n  Type de tiers:")
    print("    [1] Fournisseurs (mouvements créditeurs)")
    print("    [2] Clients (solde)")
    print("    [3] Retour")
    choix_type = input("\n  Votre choix (1-3): ").strip()
    if choix_type == '3': return
    type_tiers = 'fournisseurs' if choix_type == '1' else 'clients'
    dossier_input = DIRS['sel_input_frs'] if type_tiers == 'fournisseurs' else DIRS['sel_input_clients']

    # Choisir fichier
    fichier = choisir_fichier(dossier_input, message=f"Fichiers {type_tiers}")
    if not fichier: return

    # Charger selon le format
    print(f"\n  Chargement de {fichier.name}...")
    if est_format_dynamics(fichier):
        print("  → Format Dynamics AX détecté")
        df, societe = parser_dynamics_ax(fichier, type_tiers)
        if df is None or df.empty:
            print("  ✘ Aucun tiers trouvé dans ce fichier")
            return
        if societe: print(f"  Société: {societe}")
        print(f"  ✓ {len(df)} tiers chargés")

        # Le DataFrame est déjà structuré
        col_montant = 'total_credit' if type_tiers == 'fournisseurs' else 'solde_closing'
        col_tiers = 'code_tiers'
        col_nom = 'nom'
    else:
        print("  → Format Excel standard détecté")
        df = charger_fichier(fichier)
        if df is None or df.empty:
            print("  ✘ Fichier vide ou illisible")
            return
        print(f"  ✓ {len(df)} lignes chargées")

        # Mapping interactif
        mapping = mapper_colonnes(df, type_tiers, config)
        col_tiers = mapping.get('tiers')
        col_nom = mapping.get('nom_tiers', col_tiers)
        if type_tiers == 'fournisseurs':
            col_montant = mapping.get('montant_credit', mapping.get('solde'))
        else:
            col_montant = mapping.get('solde')

        if not col_tiers or not col_montant:
            print("  ✘ Colonnes essentielles non identifiées")
            return

        # Nettoyer et grouper
        df[col_montant] = pd.to_numeric(df[col_montant], errors='coerce').fillna(0)
        df_grouped = df.groupby(col_tiers).agg({
            col_montant: 'sum',
            **({col_nom: 'first'} if col_nom != col_tiers else {})
        }).reset_index()
        df = df_grouped
        # Renommer pour uniformiser
        df = df.rename(columns={col_tiers:'code_tiers'})
        if col_nom != col_tiers and col_nom in df.columns:
            df = df.rename(columns={col_nom:'nom'})
        else:
            df['nom'] = df['code_tiers']
        df = df.rename(columns={col_montant: 'montant'})
        col_montant = 'montant'

    # Demander intervalle
    defaut = config['seuils_selection'].get(f"{type_tiers}_intervalle_defaut", "85-95")
    intervalle = input(f"\n  Intervalle de couverture [{defaut}]: ").strip() or defaut
    try:
        parts = intervalle.replace(' ','').split('-')
        seuil_min = float(parts[0]) / 100
        seuil_max = float(parts[1]) / 100 if len(parts) > 1 else seuil_min + 0.10
    except:
        seuil_min, seuil_max = 0.85, 0.95

    # Filtrer positifs pour clients
    if type_tiers == 'clients':
        df_positifs = df[df[col_montant] > 0]
    else:
        df_positifs = df[df[col_montant].abs() > 0]

    # Sélection
    sel, stats = selectionner_par_cumul(df_positifs, col_montant, seuil_min, seuil_max)

    # Ajouter anomalies (clients créditeurs)
    if type_tiers == 'clients':
        anomalies = df[df.get('solde_closing', df.get(col_montant, pd.Series())) < 0]
        if not anomalies.empty:
            anomalies = anomalies.copy()
            anomalies['motif_selection'] = 'ANOMALIE'
            sel = pd.concat([sel, anomalies], ignore_index=True)

    # Rapport
    print(f"\n  ╔{'═'*56}╗")
    print(f"  ║  RÉSULTAT SÉLECTION — {type_tiers.upper():<33}║")
    print(f"  ╚{'═'*56}╝")
    print(f"  Population: {stats.get('total_pop',len(df))} tiers")
    print(f"  Sélectionnés: {len(sel)} (couverture: {stats.get('couverture',0):.1%})")
    print(f"  Total montant: {formater_montant(stats.get('total_montant',0))}")
    print()

    col_aff_nom = 'nom' if 'nom' in sel.columns else 'code_tiers'
    for _, r in sel.iterrows():
        nom = str(r.get(col_aff_nom, r.get('code_tiers','')))[:40]
        code = r.get('code_tiers','')
        montant = r.get(col_montant, r.get('solde_closing', 0))
        motif = r.get('motif_selection','')
        pct = r.get('pct_du_total','')
        print(f"    {code:<14} {nom:<42} {montant:>15,.2f} {motif}")

    # Export Excel
    chemin_sortie = SCRIPT_DIR / DIRS['sel_output']
    date_str = datetime.now().strftime('%Y%m%d')
    nom_export = f"selection_{type_tiers}_{fichier.stem}_{date_str}.xlsx"
    chemin_excel = chemin_sortie / nom_export

    # Ajouter colonnes email et adresse
    if 'email' not in sel.columns: sel['email'] = ''
    if 'adresse' not in sel.columns: sel['adresse'] = ''

    sel.to_excel(chemin_excel, index=False, engine='openpyxl')

    # Formater l'Excel
    from openpyxl import load_workbook
    from openpyxl.styles import Font, PatternFill, Border, Side
    wb = load_workbook(chemin_excel)
    ws = wb.active
    hdr_fill = PatternFill('solid', fgColor='2E75B6')
    hdr_font = Font(bold=True, color='FFFFFF', name='Arial', size=10)
    vert_fill = PatternFill('solid', fgColor='E2EFDA')
    border = Border(left=Side('thin'),right=Side('thin'),top=Side('thin'),bottom=Side('thin'))

    for cell in ws[1]:
        cell.fill, cell.font, cell.border = hdr_fill, hdr_font, border
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, max_col=ws.max_column):
        for cell in row:
            cell.border = border
            cell.font = Font(name='Arial', size=10)
            if cell.column_letter in ('E','F') or (ws.cell(1,cell.column).value or '').lower() in ('email','adresse'):
                if not cell.value: cell.fill = vert_fill
    ws.freeze_panes = 'A2'
    for col in ws.columns:
        ws.column_dimensions[col[0].column_letter].width = max(12, min(40, max(len(str(c.value or '')) for c in col)+2))
    wb.save(chemin_excel)

    print(f"\n  ✓ Export: {chemin_excel}")
    print(f"  → Remplissez les colonnes Email et Adresse (vertes)")
    logging.info(f"Sélection {type_tiers}: {len(sel)} tiers, couverture {stats.get('couverture',0):.1%}")
    return chemin_excel


# ═══════════════════════════════════════════════════════════════
#  PHASE 2 : GÉNÉRATION DES LETTRES
# ═══════════════════════════════════════════════════════════════

PATTERNS_TIERS = [r'«\s*.*?\s*»']
PATTERNS_MONTANT = [r'MAD\s+[\d\s]+[\d,]+[\d.]*\d+', r'EUR\s+[\d\s]+[\d,]+[\d.]*\d+']
PATTERNS_DATE = [
    r'\d{1,2}\s+(?:janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre)\s+\d{4}',
    r'\d{1,2}\s+(?:JANVIER|FÉVRIER|MARS|AVRIL|MAI|JUIN|JUILLET|AOÛT|SEPTEMBRE|OCTOBRE|NOVEMBRE|DÉCEMBRE)\s+\d{4}',
]
PATTERNS_ENTETE = [r'En\s+tête\s+de\s+la\s+société\s+auditée', r'^logo$']

def analyser_template(chemin):
    """Analyse un template Word et détecte les zones dynamiques."""
    doc = DocxDocument(chemin)
    zones, texte = [], ""
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text: continue
        texte += text + "\n"
        for pat in PATTERNS_TIERS:
            for m in re.findall(pat, text):
                zones.append({'type':'tiers_name','para':idx,'original':m})
        for pat in PATTERNS_MONTANT:
            for m in re.findall(pat, text, re.I):
                zones.append({'type':'montant','para':idx,'original':m})
        for pat in PATTERNS_DATE:
            for m in re.findall(pat, text, re.I):
                dt = 'date_cloture' if 'décembre' in m.lower() else 'date_lettre'
                if any(kw in text.lower() for kw in ['arrêté','clos','clôture','décembre 202']):
                    dt = 'date_cloture'
                zones.append({'type':dt,'para':idx,'original':m})
        for pat in PATTERNS_ENTETE:
            if re.search(pat, text, re.I):
                zones.append({'type':'entete','para':idx,'original':text})
    # Noms en dur
    fixes = {'messieurs','messieurs,','le','cachet','signature','observations'}
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if (3 < len(text) < 40 and idx < 12 and text.lower() not in fixes
            and not any(z['para']==idx for z in zones)
            and not any(kw in text.lower() for kw in ['fidaroc','boulevard','fax','tel',
                        'messieurs','commissaire','cachet','nous','compte','renseignement'])):
            zones.append({'type':'tiers_name_dur','para':idx,'original':text})
    # Dédupliquer
    seen = set()
    zones = [z for z in zones if (k:=(z['type'],z['para'],z['original'])) not in seen and not seen.add(k)]
    # Type de template par contenu puis par nom de fichier
    tl = texte.lower()
    tt = 'generique'
    if 'client' in tl or 'solde en notre faveur' in tl: tt = 'clients'
    elif 'fournisseur' in tl or 'relevé détaillé' in tl: tt = 'fournisseurs'
    elif 'assurance' in tl or 'police' in tl: tt = 'assurances'
    elif 'banque' in tl or 'bancaire' in tl: tt = 'banques'
    elif 'avocat' in tl or 'litige' in tl: tt = 'avocats'
    if tt == 'generique':
        nl = Path(chemin).name.lower()
        if 'client' in nl: tt = 'clients'
        elif 'fournisseur' in nl or 'frs' in nl: tt = 'fournisseurs'
        elif 'assurance' in nl: tt = 'assurances'
        elif 'banque' in nl or 'bank' in nl: tt = 'banques'
        elif 'avocat' in nl: tt = 'avocats'
    return {'zones':zones, 'type':tt, 'chemin':chemin}

def remplir_template(chemin_template, zones, tiers, config_cab, date_cloture):
    """Remplace les zones dynamiques dans un template Word. Retourne le Document."""
    doc = DocxDocument(chemin_template)
    nom = tiers.get('nom', tiers.get('code_tiers',''))
    solde = tiers.get('solde_closing', tiers.get('montant',0))
    devise = tiers.get('devise','MAD')
    montant_fmt = formater_montant(solde, devise)
    date_clot_fmt = formater_date_fr(date_cloture, 'long')
    date_lettre_fmt = formater_date_fr(style='lettre')
    nom_societe = config_cab.get('nom_societe', config_cab.get('nom',''))

    for z in zones:
        pi = z['para']
        if pi >= len(doc.paragraphs): continue
        para = doc.paragraphs[pi]
        orig = z['original']
        remplacement = {
            'tiers_name': nom, 'tiers_name_dur': nom,
            'montant': montant_fmt, 'date_cloture': date_clot_fmt,
            'date_lettre': date_lettre_fmt, 'entete': nom_societe,
        }.get(z['type'], '')
        if not remplacement: continue
        # Remplacement dans les runs
        for run in para.runs:
            if orig in run.text:
                run.text = run.text.replace(orig, remplacement)
                break
        else:
            if orig in para.text and para.runs:
                nouveau = para.text.replace(orig, remplacement)
                for run in para.runs: run.text = ""
                para.runs[0].text = nouveau
    return doc

def phase2_generation_lettres(config):
    """Phase interactive de génération des lettres."""
    print("\n" + "="*60)
    print("  PHASE 2 : GÉNÉRATION DES LETTRES")
    print("="*60)

    print("\n  Source des données:")
    print("    [1] À partir de la sélection (selection/output/)")
    print("    [2] À partir d'un fichier personnalisé (lettres/input/)")
    print("    [3] Retour")
    choix = input("\n  Votre choix (1-3): ").strip()
    if choix == '3': return

    if choix == '1':
        fichier = choisir_fichier(DIRS['sel_output'], ['.xlsx','.xls'], "Fichiers de sélection")
    else:
        fichier = choisir_fichier(DIRS['lettres_input'], message="Fichiers disponibles")
    if not fichier: return

    # Déterminer le type
    nl = fichier.name.lower()
    if 'fournisseur' in nl or 'frs' in nl or 'vendor' in nl:
        type_tiers = 'fournisseurs'
    elif 'client' in nl or 'customer' in nl:
        type_tiers = 'clients'
    else:
        t = input("  Type ? [1] Clients [2] Fournisseurs: ").strip()
        type_tiers = 'fournisseurs' if t == '2' else 'clients'

    # Charger les données
    df = pd.read_excel(fichier)
    print(f"  ✓ {len(df)} tiers chargés depuis {fichier.name}")

    # Trouver le template
    canvas_dir = SCRIPT_DIR / DIRS['sel_canvas']
    templates = [f for f in canvas_dir.iterdir() if f.suffix == '.docx' and not f.name.startswith('~')]
    template = None
    for t in templates:
        analyse = analyser_template(str(t))
        if analyse['type'] == type_tiers:
            template = analyse; break
    if not template:
        for t in templates:
            a = analyser_template(str(t))
            if a['type'] == 'generique':
                template = a; break
    if not template:
        print(f"\n  ✘ Aucun template trouvé pour '{type_tiers}'")
        print(f"    Placez un fichier .docx contenant '{type_tiers}' dans {canvas_dir}")
        print(f"    Fichiers trouvés: {[t.name for t in templates]}")
        return

    print(f"  ✓ Template: {Path(template['chemin']).name} ({template['type']})")
    print(f"    Zones: {len(template['zones'])}")
    for z in template['zones']:
        print(f"      [{z['type']:<18}] → \"{z['original'][:50]}\"")

    # Demander infos complémentaires
    date_clot = input(f"  Date de clôture [{config['cabinet']['date_arretes_defaut']}]: ").strip()
    date_clot = date_clot or config['cabinet']['date_arretes_defaut']
    nom_societe = input(f"  Nom de la société auditée: ").strip()
    config_gen = {**config['cabinet'], 'nom_societe': nom_societe, 'date_cloture': date_clot}

    # Identifier les colonnes du DataFrame
    col_nom = None
    for c in df.columns:
        if str(c).lower() in ('nom','name','nom_tiers','nom / raison sociale','libellé'):
            col_nom = c; break
    col_code = None
    for c in df.columns:
        if str(c).lower() in ('code_tiers','code','compte','n°','code client','code fournisseur'):
            col_code = c; break
    col_solde = None
    for c in df.columns:
        if any(kw in str(c).lower() for kw in ('solde','montant','closing','credit','balance')):
            col_solde = c; break

    # Générer les lettres
    sortie = SCRIPT_DIR / DIRS['lettres_output']
    sortie.mkdir(parents=True, exist_ok=True)
    # Vider le dossier temp
    temp_dir = SCRIPT_DIR / DIRS['lettres_temp']
    temp_dir.mkdir(parents=True, exist_ok=True)

    fichiers_gen = []
    print(f"\n  Génération de {len(df)} lettres...")
    for _, row in df.iterrows():
        tiers = {
            'nom': str(row.get(col_nom, row.get(col_code, f'Tiers_{_}'))),
            'code_tiers': str(row.get(col_code, '')),
            'solde_closing': float(row[col_solde]) if col_solde and pd.notna(row.get(col_solde)) else 0,
            'devise': str(row.get('devise','MAD')),
        }
        doc = remplir_template(template['chemin'], template['zones'], tiers, config_gen, date_clot)
        nom_safe = nettoyer_nom(tiers['nom'])
        code = tiers['code_tiers']
        nom_fic = f"LC_{code}_{nom_safe}.docx" if code else f"LC_{nom_safe}.docx"
        chemin_out = sortie / nom_fic
        doc.save(str(chemin_out))
        fichiers_gen.append({'chemin':str(chemin_out),'nom':tiers['nom'],'code':code})
        print(f"    ✓ {nom_fic}")

    # Consolider en un seul Word
    if fichiers_gen:
        nom_cons = f"Circularisation_{type_tiers.capitalize()}_{nettoyer_nom(nom_societe)}.docx"
        chemin_cons = sortie / nom_cons
        doc_final = DocxDocument(fichiers_gen[0]['chemin'])
        for fg in fichiers_gen[1:]:
            doc_final.add_page_break()
            doc_src = DocxDocument(fg['chemin'])
            for para in doc_src.paragraphs:
                new_p = doc_final.add_paragraph()
                new_p.paragraph_format.alignment = para.paragraph_format.alignment
                for run in para.runs:
                    nr = new_p.add_run(run.text)
                    nr.bold, nr.italic, nr.underline = run.bold, run.italic, run.underline
                    if run.font.size: nr.font.size = run.font.size
                    if run.font.name: nr.font.name = run.font.name
        doc_final.save(str(chemin_cons))
        print(f"\n  ✓ Consolidé: {nom_cons} ({len(fichiers_gen)} lettres)")

        # Sauvegarder le mapping pour le split
        mapping = [{'tiers':f['nom'],'code':f['code']} for f in fichiers_gen]
        etat = {'mapping_pages':mapping, 'type':type_tiers, 'date':datetime.now().isoformat()}
        with open(sortie / 'etat_generation.json','w',encoding='utf-8') as f:
            json.dump(etat, f, indent=2, ensure_ascii=False)

    print(f"\n  ✓ {len(fichiers_gen)} lettres dans {sortie}")
    logging.info(f"Génération {type_tiers}: {len(fichiers_gen)} lettres")


# ═══════════════════════════════════════════════════════════════
#  PHASE 3 : SPLIT DES LETTRES PDF
# ═══════════════════════════════════════════════════════════════

def phase3_split_pdf(config):
    """Phase interactive de split des PDF signés."""
    print("\n" + "="*60)
    print("  PHASE 3 : SPLIT DES LETTRES PDF")
    print("="*60)

    # Choisir le PDF
    fichier = choisir_fichier(DIRS['lettres_input'], ['.pdf'], "PDF à découper")
    if not fichier: return

    reader = PdfReader(str(fichier))
    total_pages = len(reader.pages)
    print(f"\n  PDF: {fichier.name} ({total_pages} pages)")

    # Charger le mapping si disponible
    etat_path = SCRIPT_DIR / DIRS['lettres_output'] / 'etat_generation.json'
    mapping = None
    if etat_path.exists():
        with open(etat_path,'r',encoding='utf-8') as f:
            etat = json.load(f)
            mapping = etat.get('mapping_pages', [])
        if mapping:
            print(f"  ✓ Mapping trouvé: {len(mapping)} tiers attendus")
            pages_est = max(1, total_pages // len(mapping))
            print(f"  → Estimation: {pages_est} page(s)/lettre")

    # Demander pages par lettre
    defaut = str(max(1, total_pages // len(mapping))) if mapping else "1"
    ppl = input(f"  Pages par lettre [{defaut}]: ").strip() or defaut
    try: pages_par_lettre = int(ppl)
    except: pages_par_lettre = 1

    # Si pas de mapping, créer un mapping générique
    if not mapping:
        nb_lettres = total_pages // pages_par_lettre
        mapping = [{'tiers':f'Lettre_{i+1}','code':f'{i+1:03d}'} for i in range(nb_lettres)]

    # Split
    sortie = SCRIPT_DIR / DIRS['lettres_output']
    sortie.mkdir(parents=True, exist_ok=True)
    page_courante = 0
    fichiers_split = []

    print(f"\n  Découpage en cours...")
    for i, tiers_info in enumerate(mapping):
        nom = tiers_info.get('tiers', f'Lettre_{i+1}')
        code = tiers_info.get('code', '')

        debut = page_courante
        fin = min(debut + pages_par_lettre, total_pages)
        if debut >= total_pages:
            print(f"  ⚠ Plus de pages pour {nom}")
            break

        writer = PdfWriter()
        for p in range(debut, fin):
            writer.add_page(reader.pages[p])

        nom_safe = nettoyer_nom(nom)
        nom_fic = f"LC_{code}_{nom_safe}.pdf" if code else f"LC_{nom_safe}.pdf"
        chemin_out = sortie / nom_fic
        with open(chemin_out, 'wb') as f:
            writer.write(f)

        fichiers_split.append({
            'chemin': str(chemin_out), 'chemin_pdf': str(chemin_out),
            'tiers': nom, 'nom': nom, 'code': code,
            'pages': f"{debut+1}-{fin}"
        })
        print(f"    ✓ {nom_fic:<55} (pages {debut+1}-{fin})")
        page_courante = fin

    restantes = total_pages - page_courante
    if restantes > 0:
        print(f"\n  ⚠ {restantes} page(s) restante(s)")
    print(f"\n  ✓ {len(fichiers_split)} lettres individuelles dans {sortie}")

    # Sauvegarder la liste pour les emails
    with open(sortie / 'split_result.json','w',encoding='utf-8') as f:
        json.dump(fichiers_split, f, indent=2, ensure_ascii=False)

    logging.info(f"Split PDF: {len(fichiers_split)} lettres depuis {fichier.name}")
    return fichiers_split

# ═══════════════════════════════════════════════════════════════
#  PHASE 4 : GÉNÉRATION DES EMAILS
# ═══════════════════════════════════════════════════════════════

DEFAULT_TEMPLATE_CLIENT = """Objet: Circularisation - Confirmation de solde client

Bonjour,

Dans le cadre de notre mission de commissariat aux comptes pour la société "{nom_societe}" au titre de l'exercice clos le {date_arretes}, nous vous prions de bien vouloir répondre à la lettre de circularisation ci-jointe.

En vous remerciant par avance pour votre retour.

Bien cordialement

{nom_cabinet}"""

DEFAULT_TEMPLATE_FRS = """Objet: Circularisation - Confirmation fournisseur

Bonjour,

Dans le cadre de notre mission de commissariat aux comptes pour la société "{nom_societe}" au titre de l'exercice clos le {date_arretes}, nous vous prions de bien vouloir répondre à la lettre de circularisation ci-jointe.

En vous remerciant par avance pour votre retour.

Bien cordialement

{nom_cabinet}"""

def creer_templates_email():
    """Crée les templates par défaut s'ils n'existent pas."""
    tpl_dir = SCRIPT_DIR / DIRS['emails_templates']
    tpl_dir.mkdir(parents=True, exist_ok=True)
    for nom, contenu in [('template_client.txt', DEFAULT_TEMPLATE_CLIENT),
                          ('template_fournisseur.txt', DEFAULT_TEMPLATE_FRS)]:
        chemin = tpl_dir / nom
        if not chemin.exists():
            with open(chemin, 'w', encoding='utf-8') as f:
                f.write(contenu)

def phase4_generation_emails(config):
    """Phase interactive de génération des emails."""
    print("\n" + "="*60)
    print("  PHASE 4 : GÉNÉRATION DES EMAILS")
    print("="*60)

    creer_templates_email()

    print("\n  Source:")
    print("    [1] À partir des lettres PDF (lettres/output/)")
    print("    [2] À partir du fichier de sélection (selection/output/)")
    print("    [3] Retour")
    choix = input("\n  Votre choix (1-3): ").strip()
    if choix == '3': return

    # Informations communes
    nom_societe = input("  Nom de la société auditée: ").strip()
    date_arretes = input(f"  Date d'arrêté [{config['cabinet']['date_arretes_defaut']}]: ").strip()
    date_arretes = date_arretes or config['cabinet']['date_arretes_defaut']
    nom_cabinet = config['cabinet']['nom']

    # Déterminer le type
    t = input("  Type ? [1] Clients [2] Fournisseurs: ").strip()
    type_tiers = 'fournisseur' if t == '2' else 'client'

    # Charger le template email
    tpl_path = SCRIPT_DIR / DIRS['emails_templates'] / f"template_{type_tiers}.txt"
    with open(tpl_path, 'r', encoding='utf-8') as f:
        template_raw = f.read()

    # Extraire sujet et corps
    lignes = template_raw.strip().split('\n')
    sujet = lignes[0].replace('Objet:','').strip() if lignes[0].startswith('Objet:') else f"Circularisation - {nom_societe}"
    corps = '\n'.join(lignes[1:]).strip()
    corps = corps.format(nom_societe=nom_societe, date_arretes=date_arretes,
                          nom_cabinet=nom_cabinet, nom_tiers='{nom_tiers}',
                          montant='{montant}', code_tiers='{code_tiers}')

    if choix == '1':
        # Mode 1: à partir des PDF dans lettres/output
        pdfs = lister_fichiers(DIRS['lettres_output'], ['.pdf'])
        if not pdfs:
            print("  ⚠ Aucun PDF trouvé dans lettres/output/")
            return

        # Chercher un fichier avec les emails
        print(f"\n  {len(pdfs)} PDF trouvés.")
        print("  Avez-vous un fichier avec les emails des tiers ?")
        fichier_emails = choisir_fichier(DIRS['sel_output'], ['.xlsx','.xls'],
                                          "Fichier avec emails (ou Entrée pour ignorer)")

        emails_map = {}
        if fichier_emails:
            df_emails = pd.read_excel(fichier_emails)
            for _, row in df_emails.iterrows():
                nom = str(row.get('nom', row.get('nom_tiers', row.get('code_tiers',''))))
                email = str(row.get('email', row.get('Email','')))
                if email and '@' in email:
                    emails_map[nettoyer_nom(nom).lower()] = email

        sortie_emails = SCRIPT_DIR / DIRS['emails_output']
        sortie_emails.mkdir(parents=True, exist_ok=True)
        generes, manquants = 0, []

        for pdf_path in pdfs:
            nom_tiers = pdf_path.stem.replace('LC_','').replace('_',' ')
            # Chercher email
            email_dest = emails_map.get(nettoyer_nom(nom_tiers).lower(), '')

            # Générer .eml
            corps_final = corps.replace('{nom_tiers}', nom_tiers)
            msg = MIMEMultipart()
            msg['Subject'] = sujet
            msg['To'] = email_dest
            msg['From'] = config['cabinet']['email_expediteur']
            msg['Date'] = formatdate(localtime=True)
            msg.attach(MIMEText(corps_final, 'plain', 'utf-8'))

            with open(pdf_path, 'rb') as f:
                pj = MIMEApplication(f.read(), _subtype='pdf')
                pj.add_header('Content-Disposition', 'attachment', filename=pdf_path.name)
                msg.attach(pj)

            eml_name = f"{pdf_path.stem}.eml"
            eml_path = sortie_emails / eml_name
            with open(eml_path, 'w', encoding='utf-8') as f:
                f.write(msg.as_string())

            status = "✓" if email_dest else "⚠ (sans email)"
            print(f"    {status} {eml_name}")
            generes += 1
            if not email_dest:
                manquants.append({'tiers':nom_tiers,'motif':'Email non trouvé'})

        # Rapport manquants
        if manquants:
            rapport = sortie_emails / 'emails_manquants.csv'
            pd.DataFrame(manquants).to_csv(rapport, index=False, encoding='utf-8-sig')
            print(f"\n  ⚠ {len(manquants)} emails manquants → {rapport.name}")

        print(f"\n  ✓ {generes} emails dans {sortie_emails}")
        logging.info(f"Emails: {generes} générés, {len(manquants)} manquants")

    elif choix == '2':
        # Mode 2: à partir du fichier de sélection
        fichier = choisir_fichier(DIRS['sel_output'], ['.xlsx'], "Fichier de sélection")
        if not fichier: return

        df = pd.read_excel(fichier)
        pdf_dir = SCRIPT_DIR / DIRS['lettres_output']
        sortie_emails = SCRIPT_DIR / DIRS['emails_output']
        sortie_emails.mkdir(parents=True, exist_ok=True)
        generes, manquants = 0, []

        for _, row in df.iterrows():
            nom = str(row.get('nom', row.get('nom_tiers', row.get('code_tiers',''))))
            code = str(row.get('code_tiers',''))
            email = str(row.get('email', row.get('Email','')))
            montant = row.get('solde_closing', row.get('montant',0))

            # Chercher le PDF
            nom_safe = nettoyer_nom(nom)
            pdf_candidates = list(pdf_dir.glob(f"*{nom_safe}*.pdf")) + list(pdf_dir.glob(f"*{code}*.pdf"))
            pdf_path = pdf_candidates[0] if pdf_candidates else None

            if not pdf_path:
                manquants.append({'tiers':nom,'montant':montant,'motif':'PDF manquant'})
                continue
            if not email or '@' not in str(email):
                manquants.append({'tiers':nom,'montant':montant,'motif':'Email manquant'})
                email = ''

            corps_final = corps.replace('{nom_tiers}',nom).replace('{montant}',formater_montant(montant)).replace('{code_tiers}',code)
            msg = MIMEMultipart()
            msg['Subject'] = sujet
            msg['To'] = email
            msg['From'] = config['cabinet']['email_expediteur']
            msg['Date'] = formatdate(localtime=True)
            msg.attach(MIMEText(corps_final, 'plain', 'utf-8'))

            with open(pdf_path,'rb') as f:
                pj = MIMEApplication(f.read(), _subtype='pdf')
                pj.add_header('Content-Disposition','attachment',filename=pdf_path.name)
                msg.attach(pj)

            eml_name = f"Email_{code}_{nom_safe}.eml"
            with open(sortie_emails / eml_name, 'w', encoding='utf-8') as f:
                f.write(msg.as_string())
            print(f"    ✓ {eml_name}")
            generes += 1

        if manquants:
            rapport = sortie_emails / 'emails_manquants.csv'
            pd.DataFrame(manquants).to_csv(rapport, index=False, encoding='utf-8-sig')
            print(f"\n  ⚠ {len(manquants)} problèmes → {rapport.name}")
        print(f"\n  ✓ {generes} emails dans {sortie_emails}")

# ═══════════════════════════════════════════════════════════════
#  DIAGNOSTIC
# ═══════════════════════════════════════════════════════════════

def mode_diagnostic():
    """Vérifie l'environnement complet."""
    print("\n" + "="*60)
    print("  MODE DIAGNOSTIC")
    print("="*60)

    # Dépendances
    print("\n  Dépendances Python:")
    for mod, pip_name in _DEPS.items():
        try:
            __import__(mod)
            print(f"    ✓ {pip_name}")
        except: print(f"    ✘ {pip_name} (manquant)")

    # Dossiers
    print("\n  Dossiers:")
    for nom, chemin in DIRS.items():
        p = SCRIPT_DIR / chemin
        fichiers = len(list(p.iterdir())) if p.exists() else 0
        status = f"✓ ({fichiers} fichier(s))" if p.exists() else "✘ (manquant)"
        print(f"    {status} {chemin}")

    # Config
    print("\n  Configuration:")
    cfg_path = SCRIPT_DIR / "config.json"
    print(f"    {'✓' if cfg_path.exists() else '✘'} config.json")

    # Templates
    print("\n  Templates Word (canvas):")
    canvas = SCRIPT_DIR / DIRS['sel_canvas']
    if canvas.exists():
        for f in canvas.iterdir():
            if f.suffix == '.docx':
                a = analyser_template(str(f))
                print(f"    ✓ {f.name} → type: {a['type']} ({len(a['zones'])} zones)")
    else:
        print(f"    ✘ Dossier canvas/ non trouvé")

    input("\n  Appuyez sur Entrée pour continuer...")

# ═══════════════════════════════════════════════════════════════
#  MENU PRINCIPAL
# ═══════════════════════════════════════════════════════════════

def afficher_etat():
    """Affiche l'état du système."""
    etats = []
    for nom, chemin in [('selection/input/clients', DIRS['sel_input_clients']),
                         ('selection/input/fournisseurs', DIRS['sel_input_frs'])]:
        p = SCRIPT_DIR / chemin
        n = len(lister_fichiers(p)) if p.exists() else 0
        if n: etats.append(f"✓ {n} fichier(s) dans {nom}")
    canvas = SCRIPT_DIR / DIRS['sel_canvas']
    n = len([f for f in canvas.iterdir() if f.suffix=='.docx']) if canvas.exists() else 0
    if n: etats.append(f"✓ {n} template(s) Word dans canvas/")
    else: etats.append("⚠ Aucun template Word dans canvas/")
    out = SCRIPT_DIR / DIRS['sel_output']
    n = len(lister_fichiers(out, ['.xlsx'])) if out.exists() else 0
    if n: etats.append(f"✓ {n} sélection(s) dans selection/output/")
    lettres = SCRIPT_DIR / DIRS['lettres_output']
    n = len(lister_fichiers(lettres, ['.docx','.pdf'])) if lettres.exists() else 0
    if n: etats.append(f"✓ {n} lettre(s) dans lettres/output/")
    return etats

def menu_principal():
    config = charger_config()
    setup_logging()
    creer_dossiers()
    creer_templates_email()

    while True:
        etats = afficher_etat()
        print(f"""
{'='*62}
    AUTOMATISATION CIRCULARISATION - AUDIT
{'='*62}

  État du système:""")
        for e in etats: print(f"    {e}")
        print(f"""
  [1] Sélection des tiers (fournisseurs/clients)
  [2] Génération des lettres de circularisation
  [3] Split des lettres PDF
  [4] Génération des emails
  [5] Mode diagnostic
  [6] Quitter
""")
        choix = input("  Votre choix (1-6): ").strip()
        if choix == '1': phase1_selection(config)
        elif choix == '2': phase2_generation_lettres(config)
        elif choix == '3': phase3_split_pdf(config)
        elif choix == '4': phase4_generation_emails(config)
        elif choix == '5': mode_diagnostic()
        elif choix == '6':
            print("\n  Au revoir !")
            break
        else:
            print("  Choix invalide")

if __name__ == "__main__":
    menu_principal()

