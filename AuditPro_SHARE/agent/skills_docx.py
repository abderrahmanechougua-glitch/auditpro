"""
DOCX Skill Implementation for AuditPro modules.

Provides Word document generation capabilities:
- SRM document creation with professional formatting
- Mail merge for circularisation letters
- Table and chart embedding
- Comment and annotation support
"""
from __future__ import annotations
from pathlib import Path
from typing import Optional, Dict, Any, List
import logging

logger = logging.getLogger(__name__)


class DOCXSkillHandler:
    """Handles Word document generation for audit modules."""
    
    def __init__(self):
        self.docx_available = True  # docx is in requirements
    
    def create_document(
        self,
        title: str,
        output_path: str,
        content: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Create a new Word document.
        
        Args:
            title: Document title
            output_path: Path to save document
            content: Optional content dict with sections
        
        Returns:
            Dict with document metadata
        """
        result = {
            "success": False,
            "document_path": str(output_path),
            "title": title,
            "elements_added": 0
        }
        
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            
            doc = Document()
            
            # Add title
            title_paragraph = doc.add_heading(title, level=1)
            title_paragraph.alignment = 1  # Center alignment
            
            result["elements_added"] += 1
            
            # Add content if provided
            if content:
                for section_title, section_data in content.items():
                    doc.add_heading(section_title, level=2)
                    result["elements_added"] += 1
                    
                    if isinstance(section_data, str):
                        doc.add_paragraph(section_data)
                        result["elements_added"] += 1
                    elif isinstance(section_data, list):
                        for item in section_data:
                            doc.add_paragraph(item, style='List Bullet')
                            result["elements_added"] += 1
                    elif isinstance(section_data, dict):
                        # Add as table
                        table = doc.add_table(rows=1, cols=len(section_data))
                        table.style = 'Light Grid Accent 1'
                        
                        # Header row
                        hdr_cells = table.rows[0].cells
                        for i, (k, v) in enumerate(section_data.items()):
                            hdr_cells[i].text = str(k)
                        
                        result["elements_added"] += 1
            
            # Save document
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            result["success"] = True
            
            logger.info(f"Created document: {output_path}")
            
        except Exception as e:
            logger.error(f"Error creating document: {e}")
            result["error"] = str(e)
        
        return result
    
    def add_table(
        self,
        document_path: str,
        data: List[List[str]] | Dict,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Add a table to an existing document.
        
        Args:
            document_path: Path to Word document
            data: Table data (list of lists or DataFrame)
            title: Optional table title
        
        Returns:
            Dict with operation result
        """
        result = {
            "success": False,
            "document_path": str(document_path)
        }
        
        try:
            from docx import Document
            import pandas as pd
            
            doc = Document(document_path)
            
            # Add title if provided
            if title:
                doc.add_heading(title, level=3)
            
            # Convert data to list of lists if needed
            if isinstance(data, pd.DataFrame):
                rows = [data.columns.tolist()] + data.values.tolist()
            elif isinstance(data, dict):
                rows = [[k, v] for k, v in data.items()]
            else:
                rows = data
            
            # Create table
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = 'Light Grid Accent 1'
            
            # Populate table
            for i, row_data in enumerate(rows):
                for j, cell_data in enumerate(row_data):
                    table.rows[i].cells[j].text = str(cell_data)
            
            doc.save(document_path)
            result["success"] = True
            result["rows"] = len(rows)
            result["cols"] = len(rows[0]) if rows else 0
            
        except Exception as e:
            logger.error(f"Error adding table: {e}")
            result["error"] = str(e)
        
        return result
    
    def add_image(
        self,
        document_path: str,
        image_path: str,
        width_inches: float = 6.0,
        caption: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Add an image to document (charts, logos, etc.).
        
        Args:
            document_path: Path to Word document
            image_path: Path to image file
            width_inches: Image width in inches
            caption: Optional image caption
        
        Returns:
            Dict with operation result
        """
        result = {
            "success": False,
            "document_path": str(document_path),
            "image_path": str(image_path)
        }
        
        try:
            from docx import Document
            from docx.shared import Inches
            
            if not Path(image_path).exists():
                result["error"] = f"Image not found: {image_path}"
                return result
            
            doc = Document(document_path)
            
            # Add image
            doc.add_picture(image_path, width=Inches(width_inches))
            
            # Add caption if provided
            if caption:
                caption_paragraph = doc.add_paragraph(caption)
                caption_paragraph.alignment = 1  # Center
                caption_run = caption_paragraph.runs[0]
                caption_run.italic = True
            
            doc.save(document_path)
            result["success"] = True
            
        except Exception as e:
            logger.error(f"Error adding image: {e}")
            result["error"] = str(e)
        
        return result
    
    def mail_merge(
        self,
        template_path: str,
        merge_data: List[Dict[str, str]],
        output_dir: str,
        naming_field: str = "name"
    ) -> Dict[str, Any]:
        """
        Generate personalized documents via mail merge (circularisation letters).
        
        Args:
            template_path: Path to Word template with merge fields
            merge_data: List of dicts with merge field values
            output_dir: Directory to save merged documents
            naming_field: Field to use for output filename
        
        Returns:
            Dict with generated documents
        """
        result = {
            "success": False,
            "output_dir": str(output_dir),
            "documents_created": [],
            "errors": []
        }
        
        try:
            from docx import Document
            from docx.oxml import OxmlElement
            
            Path(output_dir).mkdir(parents=True, exist_ok=True)
            
            if not Path(template_path).exists():
                result["error"] = f"Template not found: {template_path}"
                return result
            
            for i, data_row in enumerate(merge_data):
                # Load template
                doc = Document(template_path)
                
                # Replace merge fields in document
                for paragraph in doc.paragraphs:
                    for key, value in data_row.items():
                        if f"{{{{{key}}}}}" in paragraph.text:
                            paragraph.text = paragraph.text.replace(
                                f"{{{{{key}}}}}",
                                str(value)
                            )
                
                # Also replace in tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for key, value in data_row.items():
                                if f"{{{{{key}}}}}" in cell.text:
                                    cell.text = cell.text.replace(
                                        f"{{{{{key}}}}}",
                                        str(value)
                                    )
                
                # Generate output filename
                if naming_field in data_row:
                    safe_name = str(data_row[naming_field]).replace("/", "_").replace("\\", "_")
                    output_filename = f"{safe_name}_{i+1:03d}.docx"
                else:
                    output_filename = f"merged_{i+1:03d}.docx"
                
                output_path = Path(output_dir) / output_filename
                
                try:
                    doc.save(str(output_path))
                    result["documents_created"].append(str(output_path))
                except Exception as e:
                    result["errors"].append(f"Error saving {output_filename}: {e}")
            
            result["success"] = len(result["documents_created"]) > 0
            
        except Exception as e:
            logger.error(f"Error in mail merge: {e}")
            result["error"] = str(e)
        
        return result
    
    def add_comment(
        self,
        document_path: str,
        paragraph_index: int,
        comment_text: str,
        author: str = "Audit"
    ) -> Dict[str, Any]:
        """
        Add comment/annotation to document.
        
        Args:
            document_path: Path to Word document
            paragraph_index: Paragraph to comment on
            comment_text: Comment text
            author: Comment author
        
        Returns:
            Dict with operation result
        """
        result = {
            "success": False,
            "document_path": str(document_path)
        }
        
        try:
            from docx import Document
            
            doc = Document(document_path)
            
            if paragraph_index >= len(doc.paragraphs):
                result["error"] = f"Paragraph {paragraph_index} not found"
                return result
            
            paragraph = doc.paragraphs[paragraph_index]
            
            # Add comment indicator (docx doesn't have native comment API in python-docx)
            # So we'll add it as a footnote instead
            run = paragraph.add_run(f" [{author}: {comment_text}]")
            run.font.color.rgb = None  # Blue color for comments
            
            doc.save(document_path)
            result["success"] = True
            
        except Exception as e:
            logger.error(f"Error adding comment: {e}")
            result["error"] = str(e)
        
        return result


# Singleton instance
_docx_handler = DOCXSkillHandler()


def get_docx_skill_handler() -> DOCXSkillHandler:
    """Get the DOCX skill handler instance."""
    return _docx_handler
